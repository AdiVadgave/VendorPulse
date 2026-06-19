"""Build ZenVendorPulse_Infrastructure_and_Software_Requirements_v2_DRAFT.docx

An ANNOTATED working draft. The recommended changes are already applied:
  - §2.2 per-package list collapsed into "principal frameworks" + an OSS/SBOM
    statement (no individual libraries like uvicorn/httpx/pydantic)
  - Licence table trimmed to commercial items only (open-source needs no procurement)
  - Anthropic / Claude removed; GitHub Copilot in; Foundry-only LLM

Yellow-highlighted "EDIT HERE" notes mark every decision point so the author can
finalise the choices and then delete the notes.

Usage: python _build_infra_reqs_v2_DRAFT_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "ZenVendorPulse_Infrastructure_and_Software_Requirements_v2_DRAFT.docx"

SHELL_RED = RGBColor(0xC8, 0x10, 0x2E)
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x55, 0x55, 0x55)
EDIT = RGBColor(0xB4, 0x53, 0x09)  # amber for edit notes
HEAD_HEX = "C8102E"
ZEBRA_HEX = "F5F6F8"
BOX_HEX = "FDF3F4"


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, *, bold=False, color=None, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    parts = text.split("**")
    for idx, seg in enumerate(parts):
        if not seg:
            continue
        run = p.add_run(seg)
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        run.bold = bold or (idx % 2 == 1)
        if color is not None:
            run.font.color.rgb = color


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade(hdr[i], HEAD_HEX)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
            if r_idx % 2 == 1:
                shade(cells[i], ZEBRA_HEX)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = SHELL_RED
    r.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = INK
    r.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p


def body(doc, text, *, muted=False, size=10):
    p = doc.add_paragraph()
    parts = text.split("**")
    for idx, seg in enumerate(parts):
        if not seg:
            continue
        r = p.add_run(seg)
        r.font.size = Pt(size)
        r.font.name = "Calibri"
        r.bold = (idx % 2 == 1)
        r.font.color.rgb = MUTED if muted else INK
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    parts = text.split("**")
    for idx, seg in enumerate(parts):
        if not seg:
            continue
        r = p.add_run(seg)
        r.font.size = Pt(10)
        r.font.name = "Calibri"
        r.bold = (idx % 2 == 1)
        r.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(2)
    return p


def editnote(doc, text):
    """Yellow-highlighted 'EDIT HERE' note for the author to action then delete."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.1)
    lead = p.add_run("✎ EDIT HERE — ")
    lead.bold = True
    lead.font.size = Pt(9)
    lead.font.name = "Calibri"
    lead.font.color.rgb = EDIT
    lead.font.highlight_color = WD_COLOR_INDEX.YELLOW
    parts = text.split("**")
    for idx, seg in enumerate(parts):
        if not seg:
            continue
        r = p.add_run(seg)
        r.font.size = Pt(9)
        r.font.name = "Calibri"
        r.bold = (idx % 2 == 1)
        r.font.color.rgb = EDIT
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p


def changebox(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    shade(cell, BOX_HEX)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("What changed in this revision (v2) — and how to use this DRAFT")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.name = "Calibri"
    r.font.color.rgb = SHELL_RED
    for line in lines:
        pp = cell.add_paragraph()
        pp.paragraph_format.space_after = Pt(1)
        parts = line.split("**")
        rr = pp.add_run("•  ")
        rr.font.size = Pt(9.5)
        rr.font.name = "Calibri"
        for idx, seg in enumerate(parts):
            if not seg:
                continue
            rseg = pp.add_run(seg)
            rseg.font.size = Pt(9.5)
            rseg.font.name = "Calibri"
            rseg.bold = (idx % 2 == 1)
            rseg.font.color.rgb = INK
    doc.add_paragraph()


def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    title = doc.add_paragraph()
    tr = title.add_run("ZenVendorPulse — Infrastructure & Software Requirements")
    tr.bold = True
    tr.font.size = Pt(19)
    tr.font.color.rgb = SHELL_RED
    tr.font.name = "Calibri"
    title.paragraph_format.space_after = Pt(2)

    meta = doc.add_paragraph()
    mr = meta.add_run("Zensar for Shell   ·   Version 2.0 — WORKING DRAFT   ·   16 June 2026   ·   Confidential")
    mr.font.size = Pt(9.5)
    mr.font.color.rgb = MUTED
    mr.font.name = "Calibri"
    meta.paragraph_format.space_after = Pt(8)

    # how-to banner
    p = doc.add_paragraph()
    r = p.add_run("This is a working draft. Yellow ✎ EDIT HERE notes mark decisions for you to make. "
                  "Action each one, then delete the note before sharing.")
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.name = "Calibri"
    r.font.color.rgb = EDIT
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p.paragraph_format.space_after = Pt(8)

    changebox(doc, [
        "**LLM provider:** Anthropic / Claude removed (out of scope). Microsoft Foundry (Azure OpenAI) "
        "is the only LLM in scope.",
        "**AI pair-programming tool:** Claude Code replaced by **GitHub Copilot**.",
        "**Open-source libraries** (uvicorn, httpx, pydantic, etc.) are no longer listed individually — "
        "they need no procurement. §2.2 now states the OSS policy and points to the SBOM.",
        "**Licence table** trimmed to commercial items only.",
        "**Edge** and **datastore** now show the option pair (Front Door / App Gateway; PostgreSQL / "
        "Azure SQL) for you to choose.",
    ])

    # ---- 1. Cloud Platform ----
    h1(doc, "1. Cloud Platform & Managed Services")
    add_table(
        doc,
        ["Service", "Tier", "Role"],
        [
            ["Azure VM", "D4as v6 or D8as v6", "FastAPI + MAF backend"],
            ["Azure Static Web Apps", "Standard", "React SPA (CDN)"],
            ["Azure PostgreSQL Flexible Server  /  Azure SQL Database", "GP+HA (prod)", "Primary datastore"],
            ["Azure Key Vault", "Standard", "Secrets via Managed Identity"],
            ["Azure Front Door + WAF  /  Application Gateway + WAF", "Standard", "TLS, OWASP, origin-lock"],
            ["Azure App Insights + Azure Monitor", "Pay-as-ingested", "OTel telemetry + audit mirror"],
            ["Azure Container Registry (ACR)", "Standard", "Backend images"],
            ["Azure Blob Storage", "Standard", "Files only — transcripts / generated minutes (not a datastore)"],
            ["Microsoft Foundry / Azure OpenAI", "Responses API", "LLM + content safety"],
            ["Microsoft Graph / Entra ID", "—", "Outlook / calendar / Teams · SSO / RBAC"],
        ],
        widths=[2.7, 1.5, 2.8],
    )
    editnote(doc, "**Datastore row:** keep **PostgreSQL Flexible Server** (default) OR **Azure SQL Database** "
                  "(if that's Shell's data-platform standard) — delete the one you don't use. For cost, note the "
                  "serverless / Burstable tier. Blob is files-only; do NOT make it the datastore.")
    editnote(doc, "**Edge row:** keep **Front Door + WAF** if the app is internet-facing, OR **Application "
                  "Gateway + WAF** if all users are internal — delete the other. Confirm exposure with Shell first.")

    # ---- 2. Software ----
    h1(doc, "2. Software")

    h2(doc, "2.1 Runtimes & Package Managers")
    add_table(
        doc,
        ["Software", "Version"],
        [
            ["Python", "3.11.x"],
            ["Node.js", "20 LTS (≥ 20.19) or 22 LTS (≥ 22.12) — required by Vite 8"],
            ["npm / pip / Git", "10.x / 23+ / 2.40+"],
        ],
        widths=[2.2, 4.8],
    )

    h2(doc, "2.2 Application frameworks & open-source policy")
    body(doc, "The application is built on permissively-licensed open-source (MIT / BSD / Apache-2.0; **no "
              "copyleft / GPL components**). Exact versions are pinned in requirements.txt and package-lock.json, "
              "scanned in CI (SonarQube, Trivy, GitLeaks), and a full **SBOM is available on request**. "
              "**No open-source component requires procurement.** The principal frameworks are:")
    add_table(
        doc,
        ["Framework", "Version", "Role"],
        [
            ["FastAPI (Python)", "0.115.6", "Backend web framework / REST API"],
            ["React + Vite (TypeScript)", "19.2.4 / Vite 8", "Frontend SPA"],
            ["Microsoft Agent Framework (MAF) SDK", "≈1.8 (pin exact)", "Agent orchestration (in-process baseline)"],
            ["openai SDK", "≥ 1.50.0", "LLM client (Foundry Responses API)"],
            ["azure-ai-projects / azure-identity", "≥ 1.0.0 / ≥ 1.17.0", "Foundry client / Entra auth"],
        ],
        widths=[2.8, 1.6, 2.6],
    )
    editnote(doc, "This table replaces the old per-package list. Do **not** re-add individual libraries "
                  "(uvicorn, httpx, pydantic, etc.) — they are implementation detail covered by the SBOM, and "
                  "listing them invites procurement questions for things that cost nothing. Keep only the "
                  "stack-defining frameworks above.")

    h2(doc, "2.3 AI / Agent Stack")
    add_table(
        doc,
        ["Component", "Version / Detail"],
        [
            ["MAF SDK", "1.x (≈1.8, Jun 2026) — pin exact. In-process baseline; Foundry Hosted Agents a future option"],
            ["Microsoft Foundry", "Responses API + content safety"],
            ["Model", "gpt-4.1 / gpt-4o — **Microsoft Foundry (Azure OpenAI) only**. No third-party LLM"],
            ["Authentication", "Entra ID (DefaultAzureCredential / app-only certificate) — no API keys in code"],
        ],
        widths=[2.2, 4.8],
    )
    editnote(doc, "**Model row:** confirm the exact model — **gpt-4o** or **gpt-4.1** — and pin the dated "
                  "snapshot. Anthropic/Claude already removed; the LLMProvider abstraction stays internal only.")

    h2(doc, "2.4 Developer & CI Tooling")
    add_table(
        doc,
        ["Tool", "Version", "Purpose"],
        [
            ["GitHub Copilot", "Business / Enterprise", "AI pair-programming assistant (replaces Claude Code)"],
            ["Azure CLI (az)", "2.60+", "Azure / Foundry / Graph authentication"],
            ["Docker", "24+", "Container image builds"],
            ["ruff / eslint", "latest", "Linting in CI"],
        ],
        widths=[2.2, 1.6, 3.2],
    )
    editnote(doc, "**Tooling:** GitHub Copilot is the agreed replacement for Claude Code. If a different AI "
                  "coding tool was decided on the call, replace it here and in LIC6 below.")

    # ---- 3. Licences ----
    h1(doc, "3. Licences & Subscriptions (commercial only)")
    body(doc, "Only commercial items that require procurement are listed. Open-source components carry no "
              "licence cost and need only a licence / security review (see §2.2). Costs are indicative and "
              "depend on Shell tier and usage.")
    add_table(
        doc,
        ["#", "Licence", "Tier", "Indicative cost"],
        [
            ["LIC1", "Microsoft 365 (E3/E5) — service mailbox", "Per-user", "~$23 (E3) / ~$57 (E5) /user/mo"],
            ["LIC2", "Microsoft Entra ID", "Included in M365", "—"],
            ["LIC3", "Azure subscription (VM, SWA, datastore, Key Vault, edge, App Insights, ACR)", "Consumption", "~$450–650/mo prod; +~$80–120/mo non-prod"],
            ["LIC4", "Microsoft Foundry / Azure OpenAI (gpt-4.1 / gpt-4o)", "Pay-per-token", "~$1,000/mo — Shell.AI + TRB approval"],
            ["LIC5", "Azure DevOps / GitHub Enterprise", "SaaS", "Included in Microsoft EA"],
            ["LIC6", "GitHub Copilot", "Business / Enterprise seat", "~$19 (Business) / ~$39 (Enterprise) /user/mo"],
        ],
        widths=[0.6, 3.4, 1.5, 1.5],
    )
    editnote(doc, "Open-source components are intentionally NOT in this table — they need no procurement. "
                  "Confirm the Copilot tier (Business vs Enterprise), and replace indicative costs with "
                  "Shell-confirmed figures before sharing externally.")

    # ---- 4. Access & Identity ----
    h1(doc, "4. Access & Identity")
    add_table(
        doc,
        ["Ref", "Item", "Owner"],
        [
            ["I1–I2", "App registrations + certificate credentials", "Shell Entra ID"],
            ["I3", "Admin consent: Mail.Send, Mail.Read, Calendars.ReadWrite, OnlineMeetings.ReadWrite.All, User.Read.All", "Shell Entra ID (global admin)"],
            ["I5–I6", "Entra role groups + Application Access Policy (mailbox-scoped)", "Shell IAM / Exchange"],
        ],
        widths=[0.8, 4.4, 1.8],
    )

    # ---- 5. Assumptions ----
    h1(doc, "5. Assumptions & Constraints")
    bullet(doc, "Targets the Shell production stack (MAF + Foundry + Azure + Outlook). PoC Google components removed.")
    bullet(doc, "**The LLM provider is Microsoft Foundry (Azure OpenAI) only; Anthropic / Claude are out of scope.**")
    bullet(doc, "All open-source is permissively licensed (MIT/BSD/Apache-2.0); versions pinned and CVE-scanned in CI; SBOM available on request.")
    bullet(doc, "Secrets come from Key Vault via Managed Identity in deployed environments; .env is dev-only.")
    bullet(doc, "There is no local GPU requirement — all model inference runs remotely in Foundry.")
    editnote(doc, "Add a data-residency line if Shell requires a named region (e.g. West Europe) for Azure / "
                  "Foundry / PostgreSQL.")

    # ---- 6. Developer Workstation ----
    h1(doc, "6. Developer Workstation")
    body(doc, "No local GPU is required — all inference is remote, so a standard development laptop is sufficient.")
    add_table(
        doc,
        ["Component", "Minimum", "Recommended"],
        [
            ["CPU", "i5 11th Gen / Ryzen 5 5000 (4c/8t)", "i7/i9 12th Gen+ / Ryzen 7–9 (8c+)"],
            ["RAM", "32 GB", "32–64 GB"],
            ["Storage", "512 GB SSD, ≥ 100 GB free", "1 TB NVMe"],
            ["GPU", "None (inference is remote to Foundry)", "—"],
            ["OS", "Windows 11 Pro / macOS 13+ / Ubuntu 22.04+", "+ WSL2 on Windows"],
        ],
        widths=[1.1, 3.0, 2.9],
    )

    foot = doc.add_paragraph()
    fr = foot.add_run("ZenVendorPulse — Infrastructure & Software Requirements · Zensar for Shell · v2.0 DRAFT · 16 June 2026")
    fr.font.size = Pt(8)
    fr.font.color.rgb = MUTED
    fr.font.name = "Calibri"
    foot.paragraph_format.space_before = Pt(12)
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
