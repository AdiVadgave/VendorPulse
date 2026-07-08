"""Create a new BRD = existing BRD + two added sections (NFRs, Success Criteria).
Clones existing heading/bullet paragraphs (styles aren't resolvable by name) and
draws grid borders via XML to match the doc's existing tables. Original preserved.
"""
import copy
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "VendorPulse_BRD (1).docx"
OUT = "VendorPulse_BRD_v1.1.docx"

NAVY = RGBColor(0x1F, 0x38, 0x64)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
INK = RGBColor(0x1A, 0x1A, 0x1A)

d = Document(SRC)

H1_TMPL = next(p for p in d.paragraphs if p.style and p.style.name == "Heading 1")._p
H2_TMPL = next(p for p in d.paragraphs if p.style and p.style.name == "Heading 2")._p
BULLET_TMPL = next(p for p in d.paragraphs if p.style and p.style.name == "List Paragraph")._p


def _clone_para(template_p, text, *, run_font=None, run_size=None, run_color=None, bold=None):
    p = d.add_paragraph()  # correctly inserts before the section's sectPr
    pe = p._p
    old = pe.find(qn("w:pPr"))
    if old is not None:
        pe.remove(old)
    tmpl_pPr = template_p.find(qn("w:pPr"))
    if tmpl_pPr is not None:
        pe.insert(0, copy.deepcopy(tmpl_pPr))
    r = p.add_run(text)
    if run_font:
        r.font.name = run_font
    if run_size:
        r.font.size = Pt(run_size)
    if run_color is not None:
        r.font.color.rgb = run_color
    if bold is not None:
        r.bold = bold
    return p


def h1(text):
    return _clone_para(H1_TMPL, text)


def h2(text):
    return _clone_para(H2_TMPL, text)


def body(text):
    return d.add_paragraph(text)


def bullet(text):
    return _clone_para(BULLET_TMPL, text, run_font="Calibri", run_size=10.5, run_color=INK)


def shade(cell, hexv):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexv)
    tcPr.append(shd)


def set_cell(cell, text, *, bold=False, color=INK, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.font.color.rgb = color; r.font.name = "Calibri"


def grid_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "BFBFBF")
        borders.append(e)
    tblPr.append(borders)


def add_table(headers, rows, widths, bold_cols=()):
    t = d.add_table(rows=1, cols=len(headers))
    grid_borders(t)
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell(hdr[i], h, bold=True, color=WHITE)
        shade(hdr[i], "1F3864")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], val, bold=(i in bold_cols))
    for row in t.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)
    d.add_paragraph()
    return t


# ---------------- 11. Non-Functional Requirements ----------------
h1("11. Non-Functional Requirements")
body("The following non-functional requirements apply across all modules. Quantified targets are "
     "indicative for the MVP and will be confirmed with Shell during discovery and baselined at go-live.")
add_table(
    ["#", "Category", "Requirement"],
    [
        ["N1", "Security", "Entra ID SSO with role-based access control enforced server-side; app-to-Microsoft 365 via an app-only certificate; all secrets held in Key Vault and accessed through Managed Identity (none in code or config); TLS 1.2+ in transit; least-privilege Microsoft Graph permissions."],
        ["N2", "Edge protection", "Public entry only through Azure Application Gateway with a Web Application Firewall (OWASP rule set); the backend is not directly reachable from the internet (origin-lock)."],
        ["N3", "Data residency & privacy", "Single-tenant within Shell's Azure in an approved region; Azure AI Foundry runs inside Shell's tenant and data is not used for model training; internal-only content (e.g. private comments) is never exposed to vendors."],
        ["N4", "Auditability & traceability", "Every agent run and outbound action is logged with correlation IDs to an immutable audit trail (Application Insights / Log Analytics); records are tamper-evident and not editable."],
        ["N5", "Availability & resilience", "The production database runs in a high-availability configuration; the deterministic core degrades gracefully and remains usable if the LLM is unavailable. Target uptime ~99.5% during business hours (indicative)."],
        ["N6", "Performance", "Interactive API operations respond within ~2 seconds (p95); AI-drafted content returns within ~10 seconds; dashboards render within ~3 seconds (all indicative)."],
        ["N7", "Scalability", "Supports the quarterly vendor volume and concurrent coordinators; scales vertically (resize) or via managed Azure services without re-architecture."],
        ["N8", "Data integrity", "All figures are computed deterministically and validated (range, completeness, outliers); each cycle is stored as a locked, versioned snapshot."],
        ["N9", "Observability", "Telemetry — traces, metrics and logs — emitted via OpenTelemetry to Azure Monitor / Application Insights."],
        ["N10", "Maintainability", "Deterministic logic is decoupled from the AI layer; the LLM provider sits behind an abstraction; infrastructure is defined as code; CI runs linting, security scans and a regression suite that gates deployments."],
        ["N11", "Usability & transparency", "Browser-based single-page app on modern browsers; AI-generated content is clearly labelled (\"AI-generated — pending approval\") until approved; the UI reflects role-based permissions."],
        ["N12", "Backup & retention", "The datastore is backed up regularly; data retention and purge follow Shell's data policy (retention period to be confirmed)."],
        ["N13", "Compliance alignment", "Controls align with Shell IRM 3.492 and the EU AI Act — human oversight, deterministic figures (no fabrication), in-tenant AI, and full auditability."],
    ],
    widths=[0.5, 2.0, 4.5],
    bold_cols=(1,),
)

# ---------------- 12. Success Criteria & KPIs ----------------
h1("12. Success Criteria & Key Performance Indicators")
body("Success is measured against the business objectives in Section 3. Targets below are indicative "
     "for the MVP and will be confirmed with Shell and baselined at go-live.")
add_table(
    ["#", "Success Criterion", "How It Is Measured", "Target (indicative)"],
    [
        ["S1", "Reduced administrative effort per cycle", "Coordinator time per cycle, before vs after", "~40–50% reduction"],
        ["S2", "Faster scheduling & scorecard compilation", "Elapsed time from cycle start to scheduled meeting / compiled scorecard", "Weeks → days"],
        ["S3", "On-time scorecard submission", "% of reviewers submitting by the deadline", "≥ 90% on time"],
        ["S4", "Data consistency & quality", "% of KPIs validated deterministically; outliers flagged", "100% validated; outliers auto-flagged"],
        ["S5", "Human oversight maintained", "% of outbound actions and AI drafts approved by a human before sending", "100%"],
        ["S6", "Full auditability", "% of agent runs and outbound actions logged with correlation IDs", "100%"],
        ["S7", "No fabricated figures", "Figures computed in code; the LLM narrates only", "Zero fabricated figures"],
        ["S8", "Institutional memory in use", "Open actions carried forward; recurring issues surfaced across cycles", "Carry-forward on every cycle"],
        ["S9", "AI draft quality", "% of AI-drafted briefs / minutes accepted with only minor edits", "≥ 80%"],
        ["S10", "Adoption", "Live EGB/QBR cycles run end-to-end on VendorPulse", "≥ 1 pilot cycle at go-live, scaling thereafter"],
        ["S11", "Successful go-live", "IT-security sign-off obtained and production cutover completed", "Achieved"],
    ],
    widths=[0.5, 2.3, 2.7, 1.5],
    bold_cols=(1,),
)

h2("12.1 MVP Acceptance Criteria")
body("The MVP is considered accepted when all of the following are true:")
for b in [
    "All six modules (A–F) operate end-to-end across a live cycle.",
    "The Human-in-the-Loop approval gate is enforced on every outbound communication and key decision.",
    "All scores and metrics are computed deterministically; the LLM never computes or fabricates figures.",
    "The deterministic core continues to function with the LLM disabled.",
    "Every agent run and outbound action is recorded in an auditable, tamper-evident trail.",
    "The solution is deployed to Shell's Azure environment (single-tenant) with IT-security sign-off obtained.",
    "User Acceptance Testing is completed and signed off by the Mobility VMO team.",
]:
    bullet(b)

# ---------------- version bump ----------------
dc = d.tables[0]
for r in dc.rows:
    if r.cells[0].text.strip() == "Version":
        set_cell(r.cells[1], "1.1")
rh = d.tables[1]
nr = rh.add_row().cells
set_cell(nr[0], "1.1"); set_cell(nr[1], "Added Non-Functional Requirements and Success Criteria & KPIs sections."); set_cell(nr[2], "ZenVendorPulse Team")
for p in d.paragraphs[:8]:
    if "Version 1.0" in p.text:
        for run in p.runs:
            run.text = run.text.replace("Version 1.0", "Version 1.1")

d.save(OUT)
print("Wrote", OUT)
