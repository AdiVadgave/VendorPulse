"""Correct stale PoC-era scope in VendorPulse_PRD_Full_Questions.docx.

Writes to a NEW file (original preserved):
  - Out of Scope: remove "Real calendar/email integration (mock APIs used)"
    and "Shell IT security review or production deployment" (both now in scope).
  - In Scope: add real Microsoft Graph integration + Azure deployment bullets.
  - Module bodies: replace remaining "mock"/"in demo" references with the real
    Microsoft Graph wording (consistency with corrected scope).
Formatting, headings, bullet numbering and the clarifying questions are preserved.
"""
import copy
from docx import Document
from docx.text.paragraph import Paragraph

SRC = "VendorPulse_PRD_Full_Questions.docx"
OUT = "VendorPulse_PRD_Full_Questions_corrected.docx"

d = Document(SRC)


def replace_in_para(text_contains, find, repl):
    for p in d.paragraphs:
        if text_contains in p.text and p.runs:
            for r in p.runs:
                if find in r.text:
                    r.text = r.text.replace(find, repl)
                    return True
    return False


def delete_para(exact_text):
    for p in list(d.paragraphs):
        if p.text.strip() == exact_text:
            p._p.getparent().remove(p._p)
            return True
    return False


def clone_after(ref_para, text):
    new_p = copy.deepcopy(ref_para._p)
    ref_para._p.addnext(new_p)
    np = Paragraph(new_p, ref_para._parent)
    for r in list(np.runs):
        r._r.getparent().remove(r._r)
    np.add_run(text)
    return np


# 1) Module-body mock/demo references -> real Microsoft Graph wording
replace_in_para("one-click dispatch in demo via Outlook",
                "one-click dispatch in demo via Outlook",
                "one-click dispatch via Outlook through Microsoft Graph")
replace_in_para("which mock source", "(which mock source)", "(via Microsoft Graph)")
replace_in_para("mock email / form API",
                "(mock email / form API)",
                "(via Microsoft Graph email and a native in-app form)")
replace_in_para("mock integration; supports manual capture",
                "(mock integration; supports manual capture with AI structuring)",
                "(supports uploaded transcript or manual capture with AI structuring)")
replace_in_para("One-click distribution of minutes",
                "(mock email API)", "(via Microsoft Graph / Outlook)")

# 2) Remove the two stale Out-of-Scope bullets
delete_para("Real calendar / email integration (mock APIs used)")
delete_para("Shell IT security review or production deployment")

# 3) Add real-scope bullets to In Scope (after the Module F bullet)
modf = next(p for p in d.paragraphs if p.text.strip().startswith("Module F: Cross-Cycle"))
b1 = clone_after(modf, "Real Microsoft 365 integration via Microsoft Graph — Outlook mail, calendar, and Teams meetings")
b2 = clone_after(b1, "Deployment to Shell's Azure environment (single-tenant), with IT-security sign-off as a go-live gate")

d.save(OUT)
print("Wrote", OUT)
