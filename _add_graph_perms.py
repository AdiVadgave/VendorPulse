"""Expand the Infra & Software Requirements doc with a least-privilege
Microsoft Graph permissions justification (service account). New file; original kept.
"""
import copy
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "ZenVendorPulse_Infrastructure_and_Software_Requirements_v2 (1).docx"
OUT = "ZenVendorPulse_Infrastructure_and_Software_Requirements_v3.docx"

SHELL_RED = RGBColor(0xC8, 0x10, 0x2E)
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEAD_HEX = "C8102E"
ZEBRA_HEX = "F5F6F8"
BOX_HEX = "FDF3F4"

d = Document(SRC)


def shade(cell, hexv):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexv)
    tcPr.append(shd)


def set_cell(cell, text, *, bold=False, color=INK, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    parts = text.split("**")
    for i, seg in enumerate(parts):
        if not seg:
            continue
        r = p.add_run(seg)
        r.font.size = Pt(size); r.font.name = "Calibri"
        r.bold = bold or (i % 2 == 1)
        r.font.color.rgb = color


def grid_borders(table):
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4"); e.set(qn("w:space"), "0"); e.set(qn("w:color"), "D9D9D9")
        b.append(e)
    table._tbl.tblPr.append(b)


# ---- builders that append to body (moved into place afterward) ----
def h2(text):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = SHELL_RED; r.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(3)
    return p._p


def body(text, *, size=10, italic=False, color=INK):
    p = d.add_paragraph()
    parts = text.split("**")
    for i, seg in enumerate(parts):
        if not seg:
            continue
        r = p.add_run(seg)
        r.font.size = Pt(size); r.font.name = "Calibri"; r.bold = (i % 2 == 1); r.italic = italic; r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.05
    return p._p


def bullet(text):
    p = d.add_paragraph(style="List Bullet")
    parts = text.split("**")
    for i, seg in enumerate(parts):
        if not seg:
            continue
        r = p.add_run(seg)
        r.font.size = Pt(10); r.font.name = "Calibri"; r.bold = (i % 2 == 1); r.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(2)
    return p._p


def callout(title, text):
    t = d.add_table(rows=1, cols=1)
    grid_borders(t)
    c = t.rows[0].cells[0]; shade(c, BOX_HEX); c.text = ""
    p = c.paragraphs[0]
    r = p.add_run(title); r.bold = True; r.font.size = Pt(10.5); r.font.name = "Calibri"; r.font.color.rgb = SHELL_RED
    p2 = c.add_paragraph()
    parts = text.split("**")
    for i, seg in enumerate(parts):
        if not seg:
            continue
        rr = p2.add_run(seg); rr.font.size = Pt(10); rr.font.name = "Calibri"; rr.bold = (i % 2 == 1); rr.font.color.rgb = INK
    return t._tbl


def perm_table():
    headers = ["Permission (application, app-only)", "What VendorPulse uses it for",
               "Why this is the least privilege", "How its breadth is scoped down"]
    rows = [
        ["Calendars.ReadWrite",
         "Read attendees' free/busy (via getSchedule) to find slots; create / update / cancel the QBR meeting on the service mailbox; read RSVP responses on that event.",
         "Booking the meeting needs write, and ReadWrite already covers the free/busy read — so no separate Calendars.Read is requested. getSchedule returns only free/busy, not event contents.",
         "Application Access Policy / RBAC for Applications limited to a mail-enabled security group (service mailbox + room + the named QBR attendees) — not all mailboxes. With RBAC, write is limited to the service mailbox and only read over the attendee group."],
        ["OnlineMeetings.ReadWrite.All",
         "Generate the Teams join link when the meeting is created by the service account (app-only).",
         "App-only event creation needs this to reliably produce the Teams meeting; Graph has no narrower scope for creating online meetings.",
         "Application Access Policy limited to the single organiser / service account that hosts the meetings."],
        ["Mail.Send",
         "Send scorecard requests, reminders and the minutes from the VendorPulse service mailbox.",
         "Send is the only mail action needed — the app never reads mail (Mail.Read was removed).",
         "RBAC for Applications / Application Access Policy limited to the single service mailbox — the app cannot send from any other mailbox."],
        ["User.ReadBasic.All",
         "Resolve and validate attendees (display name + email) from the directory when refreshing the attendee list.",
         "Reduced from User.Read.All — only basic fields are needed, not full profiles. Graph has no 'read specific users' scope, so any directory read is tenant-wide by type.",
         "Returns basic fields only. Can be dropped entirely if coordinators always supply attendee email addresses directly."],
        ["OnlineMeetingTranscript.Read.All  (conditional — Option C only; see §4.2)",
         "Automatically retrieve the Teams transcript after a meeting ends, to draft the minutes — no bot joins the call.",
         "Read-only, transcript content only. Requested ONLY if automated transcript retrieval (Option C) is adopted; not needed for manual or uploaded-transcript capture.",
         "Scoped via Application Access Policy to the VMO service mailbox as meeting organiser — reads only its own organised meetings' transcripts, not the organisation's. Requires transcription to have been enabled (participant notice shown)."],
    ]
    t = d.add_table(rows=1, cols=4)
    grid_borders(t)
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, bold=True, color=WHITE)
        shade(t.rows[0].cells[i], HEAD_HEX)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, v in enumerate(row):
            set_cell(cells[ci], v, bold=(ci == 0))
            if ri % 2 == 1:
                shade(cells[ci], ZEBRA_HEX)
    widths = [1.5, 1.9, 1.9, 1.9]
    for row in t.rows:
        for ci, w in enumerate(widths):
            row.cells[ci].width = Inches(w)
    return t._tbl


# ---- build the section elements in order ----
elems = []
elems.append(h2("4.1 Microsoft Graph permissions (service account) — justification & least privilege"))
elems.append(body(
    "All Microsoft Graph permissions below are granted to the **VendorPulse service account** — a dedicated "
    "Entra ID app registration / service principal that authenticates **app-only with a certificate held in "
    "Key Vault**. They are not granted to end users, and the app never impersonates a user."))
elems.append(body(
    "Because the service runs app-only (there is no signed-in user), Graph expresses these as **application "
    "permissions, which are tenant-wide by design** — that is what the \".All\" / \"all mailboxes\" wording "
    "reflects. It denotes the permission **type**, not the breadth we actually use. We constrain the real access "
    "two ways: (1) we request the **lowest-privilege variant** that does the job, and (2) we **scope each "
    "permission to a specific set of mailboxes** with an Exchange **Application Access Policy** (or the newer "
    "**RBAC for Applications**) tied to a mail-enabled security group containing only the VMO service mailbox and "
    "the relevant QBR attendees — never the whole organisation."))
elems.append(perm_table())
elems.append(callout(
    "“Why does it have to be ‘.All’?” — the direct answer",
    "Graph has no permission that reads “just these specific other users / mailboxes.” For an app-only "
    "service account, access to anything other than ‘self’ is expressed tenant-wide (‘.All’ / "
    "‘all mailboxes’). So we keep real exposure minimal by (a) choosing the least-privilege variant "
    "(ReadBasic over Read; free/busy-only getSchedule; no Mail.Read), and (b) **scoping the service account to a "
    "defined mailbox group via Application Access Policy / RBAC for Applications** — so in practice it can touch "
    "only the VMO mailbox and the named QBR attendees, not the organisation. Calendar availability uses "
    "**getSchedule**, which returns only busy/free/tentative blocks, not meeting contents."))
elems.append(body("**Reduced or removed from the earlier request (applying least privilege):**"))
elems.append(bullet("**Mail.Read — removed.** Scorecards are collected through the native in-app form (not by "
                    "parsing email replies), and RSVPs are read from the calendar event — so the app never "
                    "needs to read any mailbox."))
elems.append(bullet("**User.Read.All → User.ReadBasic.All.** Attendee lookup needs only display name and email, "
                    "not full user profiles."))
elems.append(bullet("**Net effect:** from 5 broad permissions to 4, two of them narrowed, and every one "
                    "mailbox-scoped via Application Access Policy / RBAC for Applications."))
elems.append(body(
    "**Technical note (app-only correctness):** Microsoft's findMeetingTimes API is **delegated-only** (it requires "
    "a signed-in user), so the app-only service account uses **getSchedule** for availability instead. Its "
    "least-privilege application permission is Calendars.Read, which is already covered by the Calendars.ReadWrite "
    "requested above.", size=9.5, italic=True, color=MUTED))

# ---- 4.2 Meeting capture & bot participation ----
elems.append(h2("4.2 Meeting capture & bot participation"))
elems.append(body(
    "**Policy position: no AI bot joins meetings.** VendorPulse never sends a bot or service account into the live "
    "call — no media is captured and nothing appears in the meeting roster. Minutes are produced from a transcript "
    "that Teams itself captures, or from notes the coordinator enters. Three options, in increasing automation:"))
elems.append(bullet(
    "**Option A — Manual capture (lowest).** The coordinator pastes notes; the AI Service only structures them into "
    "minutes. No recording and no Graph permission."))
elems.append(bullet(
    "**Option B — Human-enabled transcription, uploaded (MVP default).** The organiser turns on Teams transcription "
    "(participants see the on-screen notice); the transcript is exported and uploaded to VendorPulse, which drafts "
    "the minutes. No bot and no extra Graph permission."))
elems.append(bullet(
    "**Option C — Automated transcript retrieval (convenience upgrade).** The VendorPulse service mailbox is the "
    "meeting organiser; after the meeting the service account automatically retrieves the transcript via Microsoft "
    "Graph (optionally triggered by a change-notification when the transcript is ready) and drafts the minutes. "
    "Still nothing joins the call. Requires OnlineMeetingTranscript.Read.All (scoped — see §4.1) and transcription "
    "to have been enabled."))
elems.append(callout(
    "Why this satisfies “no bot may join”",
    "Being on the invite list is not joining, and is not what grants transcript access. Access is an "
    "**organiser-scoped Graph permission reading a transcript file the meeting already produced** — a backend API "
    "call after the meeting, not a participant in it. VendorPulse captures no media and never appears in the "
    "meeting roster."))
elems.append(body(
    "In every option, capturing meeting content still requires the usual confirmations from IRM / Legal: "
    "transcription consent and on-screen notice (especially for external vendor attendees), data classification and "
    "retention of the transcript and minutes, and the EU AI Act notice that AI drafts the minutes.",
    size=9.5, italic=True, color=MUTED))

# ---- insert before "5. Assumptions & Constraints" ----
sec5 = next(p for p in d.paragraphs if p.text.strip().startswith("5. Assumptions"))
ref = sec5._p
for el in elems:
    ref.addprevious(el)

# ---- update Table 6 (Access & Identity) I3 row ----
for t in d.tables:
    for r in t.rows:
        if r.cells[0].text.strip() == "I3":
            set_cell(r.cells[1],
                     "Admin consent for application (app-only) Graph permissions: Calendars.ReadWrite, "
                     "OnlineMeetings.ReadWrite.All, Mail.Send, User.ReadBasic.All — each least-privilege and "
                     "mailbox-scoped (see §4.1). Mail.Read removed; User.Read.All reduced to User.ReadBasic.All.",
                     size=9.5)

d.save(OUT)
print("Wrote", OUT)
