"""
Update the existing EGB Proposal (Brief) Word doc with richer timeline detail
sourced from yesterday's 3-week plan, fitted to the brief's 4-week framework.

Only touches the timeline table (Table 2). Other sections left intact.

Run:
    python docs/updated/_update_egb_proposal.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# ---------------------------------------------------------------------------
# Theme (matches the pleasant palette used in the other Word docs)
# ---------------------------------------------------------------------------
C_NAVY        = RGBColor(0x1B, 0x3A, 0x6F)
C_BLUE        = RGBColor(0x0F, 0x5B, 0xA8)
C_BODY        = RGBColor(0x1E, 0x29, 0x3B)
F_HEAD_FILL   = "DCE7F4"     # soft blue header fill (used on existing header row already)
F_ALT_FILL    = "F4F7FB"     # very light blue tint for alt rows

FONT_BODY = "Calibri"

DOC_PATH = Path(__file__).resolve().parent / "word" / "Zensar _ CR 9 EGB - Proposal (Brief).docx"


# ---------------------------------------------------------------------------
# New timeline content — sourced from yesterday's 3-week day-by-day plan,
# adapted to the 4-week framework (Week 3 = UAT + release candidate only;
# Week 4 = deployment + cutover + pilot kickoff + warranty start).
# ---------------------------------------------------------------------------

PRE_MOB_OUTCOME = [
    "Shell provisions Entra ID app registrations (NonProd + Prod) and admin consent on Graph application permissions",
    "Service mailbox vendorpulse-svc@shell.com provisioned by Shell with Application Access Policy scoped to the app",
    "Azure infrastructure provisioned by Shell — Postgres Flexible Server, Key Vault, App Service Plan, App Insights / Log Analytics, Container Registry, Front Door",
    "LLM provider chosen and contract established by Shell (Anthropic or Azure OpenAI); API key deposited in Key Vault",
    "DNS hostnames and TLS certificates provisioned by Shell for Prod and NonProd",
    "Source repository and CI/CD pipeline provisioned by Shell with Zensar engineers added; service connection for deployment configured",
    "Scorecard taxonomy drafted (locked at Day-2 checkpoint); 3 named UAT coordinators identified; one pilot vendor pre-agreed",
]

WEEK1_OUTCOME = [
    "Day 1 — Joint kick-off; pack walkthrough; day-zero prerequisites confirmed; ADRs locked; clean shell-prod branch with git-history scrub; CI lint + typecheck + unit tests on green",
    "Day 2 (AM) — JSON-to-Postgres migration via Alembic; Key Vault integration via managed identity; pydantic-settings + azure-identity wired",
    "Day 2 (PM) — Design Alignment Checkpoint with Shell VMO + IT Architecture + IT Security; 12 topics decided (incl. scorecard mode); Design Decision Log committed",
    "Day 3 — MSAL OIDC end-user authentication + session cookie + role guards; Graph app-only certificate authentication working; Module B path validated by working spike",
    "Day 4 — Gmail / Google Forms / google-auth code paths fully removed; non-prod App Service running containerised build via CI/CD pipeline",
    "Day 5 — Module A end-to-end against real Graph (findMeetingTimes → event creation with Teams URL → invite via sendMail); Week 1 demo; design frozen Friday EoD",
]

WEEK2_OUTCOME = [
    "Day 6 — Module B core flow built in the chosen mode (in-app form OR Excel attachment per Day-2 decision); end-to-end stakeholder round-trip exercised",
    "Day 7 — Module B completion (validation + reminder cadence + submission tracker); Modules C, D and F ported from POC with audit writes to external_calls",
    "Day 8 — Module E live (live capture + transcript parsing + minutes generation); Outlook-friendly email templates redesigned and cross-client tested (desktop / OWA / mobile)",
    "Day 9 — Admin module complete (vendor master, roles view, LLM budget panel, audit log viewer, system health); Shell branding applied across UI and emails",
    "Day 10 — Hardening pass (per-cycle token budget, rate-limit middleware, error mapping, correlation IDs); full-cycle demo end-to-end; code freeze declared",
]

WEEK3_OUTCOME = [
    "Day 11 — UAT kick-off with 3 named Shell VMO coordinators; defect tracker established with P1/P2/P3 triage; observability finalisation begins",
    "Day 12 — UAT day 2; P1 defects deployed same-day; App Insights operator workbook live; alert rules configured (app down, agent failure rate, Graph auth failure, LLM budget)",
    "Day 13 — UAT day 3 with defect batch; regression test pass against full cycle; runbook v0.9 drafted (daily-ops checklist, P1/P2/P3 procedures, common-symptom reference)",
    "Day 14 — Final UAT day; runbook walkthrough rehearsal with Shell IT Operations; pre-cutover health checks; on-call escalation routing tested",
    "Day 15 — Release candidate baked, tagged and signed off; UAT defect log closed (P1/P2 fixed; P3 deferred to warranty); end-of-Week-3 review with Shell sponsor — NO GO-LIVE THIS WEEK",
]

WEEK4_OUTCOME = [
    "Day 16 — Bicep deployment to production resource group; Alembic migrations applied to prod Postgres; infrastructure smoke tests (App Service ↔ Postgres ↔ Key Vault); Shell IT Security final permissions review",
    "Day 17 — Release-candidate image tagged prod-<date>-<sha> and deployed to prod slot:staging; smoke tests against staging slot; Shell IT Security formal sign-off; CAB ticket filed with runbook attached",
    "Day 18 — CAB approval; DNS cutover (Shell DNS → Front Door endpoint); zero-downtime slot swap (production ← staging); 60-min observation window; pilot vendor cycle kick-off",
    "Day 19 — Pilot cycle progresses through Module A → Module B; Zensar Tech Lead pair-shadowing the lead Shell VMO coordinator; App Insights monitored live; daily check-in",
    "Day 20 — Pilot cycle continues; handover artefacts finalised (runbook v1.0, decision log, architecture pack); coordinator training session #1 delivered and recorded; end-of-engagement retro; defect-warranty period begins Monday Week 5",
]


PRE_MOB_ASK = "Refer Section 6 — every item above is a Shell-side dependency that must be in place before Day 1"
WEEK1_ASK = "Day-2 attendees: Shell VMO Product Owner, IT Architecture liaison, IT Security liaison, Brand / Comms rep"
WEEK2_ASK = "Brand sign-off on email templates by Day 8; Module B path decision honoured through code freeze"
WEEK3_ASK = "3 named Shell VMO coordinators available for UAT (Days 11–15); Shell IT Operations available for runbook walkthrough on Day 14"
WEEK4_ASK = "Shell IT Security sign-off (Day 17), Shell CAB approval (Day 18), pilot-vendor cycle pre-agreed, coordinator availability for training (Day 20)"


# ---------------------------------------------------------------------------
# Cell writers
# ---------------------------------------------------------------------------

def shd(color_hex: str):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), color_hex)
    return el


def set_cell_shading(cell, color_hex: str):
    cell._tc.get_or_add_tcPr().append(shd(color_hex))


def write_bulleted_cell(cell, items: list[str], color=C_BODY, font_size=10):
    """Replace cell content with one bulleted paragraph per item."""
    # Clear existing paragraphs (keep the underlying tc element)
    tc = cell._tc
    for p in list(cell.paragraphs):
        tc.remove(p._p)

    for idx, item in enumerate(items):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        # Bullet glyph
        bullet = p.add_run("• ")
        bullet.font.name = FONT_BODY
        bullet.font.size = Pt(font_size)
        bullet.font.color.rgb = C_BLUE
        bullet.font.bold = True
        # Item text
        run = p.add_run(item)
        run.font.name = FONT_BODY
        run.font.size = Pt(font_size)
        run.font.color.rgb = color


def write_plain_cell(cell, text: str, color=C_BODY, font_size=10, bold=False):
    """Replace cell content with a single paragraph of plain text."""
    tc = cell._tc
    for p in list(cell.paragraphs):
        tc.remove(p._p)

    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    run.font.bold = bold


def write_phase_header(cell, lines: list[str]):
    """
    Phase column cell: render the phase name with its sub-line(s) (e.g. the
    'Foundations, Migration & Design Alignment' under 'Week 1') in a tidy way.
    """
    tc = cell._tc
    for p in list(cell.paragraphs):
        tc.remove(p._p)

    for idx, line in enumerate(lines):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line)
        run.font.name = FONT_BODY
        run.font.size = Pt(11)
        run.font.color.rgb = C_NAVY if idx == 0 else C_BODY
        run.font.bold = (idx == 0)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    if not DOC_PATH.exists():
        raise SystemExit(f"Doc not found: {DOC_PATH}")

    doc = Document(DOC_PATH)

    # Find the timeline table (the one whose header row contains 'Phase' / 'Duration' / 'Key Outcome')
    timeline_tbl = None
    for tbl in doc.tables:
        header_cells = [c.text.strip().lower() for c in tbl.rows[0].cells]
        if "phase" in header_cells and "duration" in header_cells and any("outcome" in h for h in header_cells):
            timeline_tbl = tbl
            break

    if timeline_tbl is None:
        raise SystemExit("Could not find the timeline table (looking for headers Phase / Duration / Key Outcome).")

    if len(timeline_tbl.rows) < 6:
        raise SystemExit(f"Timeline table has {len(timeline_tbl.rows)} rows; expected at least 6.")

    rows = timeline_tbl.rows
    # rows[0] = header (Phase / Duration / Key Outcome / Pre-Requisites)
    # rows[1] = Pre-mobilization
    # rows[2] = Week 1
    # rows[3] = Week 2
    # rows[4] = Week 3
    # rows[5] = Week 4

    n_cols = len(rows[0].cells)
    phase_col, dur_col, outcome_col, ask_col = 0, 1, 2, 3 if n_cols >= 4 else None

    # Phase column — refresh with bold styling
    write_phase_header(rows[1].cells[phase_col], ["Pre-mobilisation"])
    write_phase_header(rows[2].cells[phase_col], ["Week 1", "Foundations, Migration & Design Alignment"])
    write_phase_header(rows[3].cells[phase_col], ["Week 2", "Development & Module Testing"])
    write_phase_header(rows[4].cells[phase_col], ["Week 3", "Hardening, Stabilisation & UAT"])
    write_phase_header(rows[5].cells[phase_col], ["Week 4", "Deployment, Pilot Go-Live & Warranty Start"])

    # Duration column — keep concise
    write_plain_cell(rows[1].cells[dur_col], "Before Day 1", bold=True)
    write_plain_cell(rows[2].cells[dur_col], "Days 1–5", bold=True)
    write_plain_cell(rows[3].cells[dur_col], "Days 6–10", bold=True)
    write_plain_cell(rows[4].cells[dur_col], "Days 11–15", bold=True)
    write_plain_cell(rows[5].cells[dur_col], "Days 16–20", bold=True)

    # Key Outcome column — bulleted day-by-day detail from yesterday's plan
    write_bulleted_cell(rows[1].cells[outcome_col], PRE_MOB_OUTCOME)
    write_bulleted_cell(rows[2].cells[outcome_col], WEEK1_OUTCOME)
    write_bulleted_cell(rows[3].cells[outcome_col], WEEK2_OUTCOME)
    write_bulleted_cell(rows[4].cells[outcome_col], WEEK3_OUTCOME)
    write_bulleted_cell(rows[5].cells[outcome_col], WEEK4_OUTCOME)

    # Pre-Requisites / Ask column (4th column) — refresh
    if ask_col is not None:
        write_plain_cell(rows[1].cells[ask_col], PRE_MOB_ASK)
        write_plain_cell(rows[2].cells[ask_col], WEEK1_ASK)
        write_plain_cell(rows[3].cells[ask_col], WEEK2_ASK)
        write_plain_cell(rows[4].cells[ask_col], WEEK3_ASK)
        write_plain_cell(rows[5].cells[ask_col], WEEK4_ASK)

    # Alternate-row tint to improve readability
    set_cell_shading(rows[2].cells[phase_col], F_ALT_FILL)
    set_cell_shading(rows[2].cells[dur_col], F_ALT_FILL)
    set_cell_shading(rows[2].cells[outcome_col], F_ALT_FILL)
    if ask_col is not None:
        set_cell_shading(rows[2].cells[ask_col], F_ALT_FILL)

    set_cell_shading(rows[4].cells[phase_col], F_ALT_FILL)
    set_cell_shading(rows[4].cells[dur_col], F_ALT_FILL)
    set_cell_shading(rows[4].cells[outcome_col], F_ALT_FILL)
    if ask_col is not None:
        set_cell_shading(rows[4].cells[ask_col], F_ALT_FILL)

    # Header row — keep header fill consistent
    for c in rows[0].cells:
        set_cell_shading(c, F_HEAD_FILL)

    doc.save(DOC_PATH)
    print(f"Updated timeline table in: {DOC_PATH.name}")


if __name__ == "__main__":
    main()
