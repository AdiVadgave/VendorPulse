"""Insert answers under each clarifying question in the corrected PRD.

- Confident answers (grounded in the codebase / agreed architecture) written plainly.
- Lower-confidence / business-policy items get a yellow [NEEDS YOUR INPUT] flag.
Updates the same corrected doc in place.
"""
import copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

DOC = "VendorPulse_PRD_Full_Questions_corrected.docx"

TEAL = RGBColor(0x0B, 0x6E, 0x63)
INK = RGBColor(0x1A, 0x1A, 0x1A)
FLAG = RGBColor(0xB4, 0x53, 0x09)

d = Document(DOC)
paras = list(d.paragraphs)

# index -> (answer, review_note_or_None)
A = {
 30: ("Seeded from the previous cycle's attendee list, held in the attendee repository (JSON today, PostgreSQL in production), keyed by cycle.", None),
 31: ("Back to the same attendee repository, as an updated per-cycle attendee record.", None),
 32: ("Once per cycle, on cycle trigger — it isn't a continuous sync.", None),
 38: ("We don't auto-resolve conflicting nominations; both are surfaced to the VMO coordinator, who decides (HITL).", "confirm this reconciliation rule"),
 39: ("Email format is validated, and for Shell users we confirm identity via the Entra/Graph directory lookup (User.Read.All). External (vendor) contacts are accepted as free-text email and aren't directory-validated.", None),
 42: ("Configurable; default 48 hours (as noted in the body).", None),
 43: ("The VMO coordinator (the cycle organiser).", None),
 46: ("Generated deterministically from configurable working hours (default 09:00–17:00) and attendee time zones; in production, real free/busy comes from Microsoft Graph findMeetingTimes. Up to ~12 candidate slots are scored and the top 3 presented.", "confirm the date window length, e.g. next 2–3 weeks"),
 48: ("On the attendee record as a role / 'key' flag in the attendee repository (roles include organiser, executive sponsor, etc.).", None),
 49: ("It escalates to the coordinator (HITL): the best-available options are shown with the gaps highlighted, and the coordinator relaxes a constraint or opens a new window.", None),
 51: ("Today the ranking is deterministic — a hard constraint (organiser + executive sponsor must be free) plus a soft score from attendance %, a conflict penalty and a time-zone bonus; ties break on the higher attendance score. Role-weighted voting is not implemented yet.", "confirm whether to add role-weighted scoring as you noted"),
 53: ("From attendee calendars — the availability service in the PoC, and Microsoft Graph free/busy in production.", None),
 54: ("Correct — no. We flag only that a conflict exists; private meeting details are never exposed (privacy).", None),
 59: ("In a configurable template held with the cycle/app configuration; production uses standard Outlook templates.", None),
 62: ("From Microsoft Graph — the event's per-attendee response status.", None),
 63: ("On dashboard load and on a scheduled background poll; the interval is configurable.", "confirm the poll interval"),
 65: ("Default: two automated nudges, then escalation.", "confirm the number of nudges"),
 66: ("Escalates to the coordinator (and the sponsor if needed) for a manual decision.", None),
 67: ("Sending invites and reminders is gated — the coordinator approves before anything goes out.", None),
 72: ("Stored per cycle as a versioned snapshot in the datastore (PostgreSQL in production); each cycle keeps its own version.", None),
 73: ("Production uses a native in-app form (not Microsoft or Google Forms), distributed via a Graph email magic-link. Collection sits behind an interface, so if Shell mandates another method we can swap the adapter without changing the rest of the system.", None),
 75: ("The VMO coordinator / organiser, enforced by RBAC.", None),
 77: ("Recommendation: editable until first submission, then locked and versioned for the cycle.", "confirm this locking policy"),
 82: ("Saved against respondent + cycle in the datastore until submission or cycle close.", "confirm the draft retention period"),
 84: ("Configured per cycle (default 48 hrs); the nudge goes to the respondent with the coordinator copied.", None),
 85: ("Escalates to the coordinator/sponsor; any still-missing input is marked in the compiled scorecard.", None),
 89: ("Cross-source discrepancies are configurable; the outlier/anomaly check uses a z-score of 1.5 SD.", "confirm the cross-source discrepancy threshold value"),
 91: ("From Module F cross-cycle memory — the prior versioned scorecards.", None),
 94: ("In the datastore, as a locked, versioned snapshot for the cycle.", None),
 98: ("Via the shared datastore (read by Modules C and F) over the internal API — not a file hand-off.", None),
 103: ("The compiled, locked scorecard from Module B (datastore).", None),
 105: ("From Module F (cross-cycle memory) and the shared action register.", None),
 107: ("Azure AI Foundry GPT-4o, inside Shell's tenant, via the AI Service. All figures are computed deterministically and passed to the model as fixed inputs; the LLM only narrates — it never computes — and dates/IDs are stamped in code.", None),
 108: ("No. Foundry runs inside Shell's tenant; the data stays in-tenant and is not used for model training.", None),
 112: ("Agenda ordering is rule-based (driven by Red/Amber KPIs and risk); the discussion-point suggestions are LLM-drafted and human-approved.", None),
 114: ("In the datastore, against the cycle.", None),
 117: ("In the datastore, as the cycle's alignment record.", None),
 119: ("It's intended as one consolidated action register shared across C/D/E and carried into F.", "confirm a single consolidated register — today actions are logged per module"),
 122: ("In the datastore; structured JSON, rendered in-app and exportable.", None),
 124: ("Through the shared datastore to Module D.", None),
 125: ("Yes — confirmed. Sponsor/organiser sign-off is required before any vendor engagement.", None),
 133: ("Currently LLM-drafted (advisory) and grounded in the scorecard evidence; they can also be seeded from Module F history.", "confirm whether to seed counters from prior-cycle history"),
 135: ("The agreed Module C positions are passed to the draft as fixed context, and the human review step confirms alignment before use.", None),
 137: ("Flagged and routed via the coordinator to Shell's legal/contracts owner; the system never interprets legal content.", "confirm the named legal owner / hand-off route"),
 140: ("Both — manual entry or an uploaded transcript parsed by the AI Service.", None),
 145: ("In the datastore, as the vendor-prep record.", None),
 147: ("Through the shared datastore to Module E.", None),
 148: ("Yes — confirmed. The organiser reviews the vendor-prep summary before the governance meeting.", None),
 153: ("Reads the compiled scorecard (B), alignment summary (C) and vendor-prep summary (D) from the datastore; the assembled pack is stored against the cycle.", None),
 155: ("An in-app structured view, exportable to PDF/DOCX. Slide auto-population stays out of scope.", "confirm the required export formats"),
 158: ("Both are supported — an uploaded transcript or manual notes.", None),
 160: ("In the meeting record (datastore), structured into the five note types: question, objection, decision, appreciation, action.", None),
 164: ("From the live-capture log / parsed notes; minutes are drafted by the AI Service (Foundry GPT-4o) and human-approved.", None),
 166: ("The consolidated, shared action register.", "confirm single register (same point as Module C)"),
 168: ("Yes — the refreshed Module A attendee list.", None),
 169: ("Yes — confirmed. The organiser reviews and edits the draft minutes before distribution.", None),
 172: ("In the datastore action register, updated through the app.", None),
 175: ("Open actions and unresolved items are carried to Module F via the shared datastore.", None),
 180: ("PostgreSQL is the system of record, holding versioned per-cycle snapshots of scorecards, decisions, actions and minutes.", "retention period is a Shell data-policy decision"),
 182: ("By structured query — vendor, KPI, cycle and time range. Free-text keyword search is a possible later enhancement.", None),
 186: ("From the cross-cycle memory store, across all available prior cycles (look-back is configurable).", None),
 188: ("Deterministically from the recent score delta: improving if the change is ≥ +0.5, declining if ≤ −0.5, otherwise stable.", None),
 190: ("Today by KPI/category tagging — a recurring issue is a category scoring below 3.0 for two or more consecutive cycles. Text-similarity matching is a possible enhancement.", "confirm tagging is sufficient vs text similarity"),
 192: ("Yes — from the shared action register populated in Module E.", None),
 196: ("A React 19 single-page app using Recharts; it's interactive — with filters and drill-down.", None),
 198: ("From the cross-cycle memory store, back to the underlying cycle scorecard.", None),
 202: ("Foundry GPT-4o, fed only the deterministically computed values — narrator only.", None),
 204: ("Today it's rule-based (recurring issue ≥ 2 cycles below 3.0; declining trajectory).", "specific risk thresholds need business sign-off"),
 205: ("Recommendation: dashboard-only for the MVP, with proactive notifications as a later enhancement.", "confirm preference"),
}


def add_answer(qpara, answer, note):
    new_p = copy.deepcopy(qpara._p)
    # strip bullet/numbering so the answer isn't another bullet
    pPr = new_p.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)
    qpara._p.addnext(new_p)
    np = Paragraph(new_p, qpara._parent)
    for r in list(np.runs):
        r._r.getparent().remove(r._r)
    np.paragraph_format.space_after = Pt(6)
    lead = np.add_run("A:  ")
    lead.bold = True; lead.italic = False; lead.font.color.rgb = TEAL; lead.font.size = Pt(10.5); lead.font.name = "Calibri"
    body = np.add_run(answer)
    body.bold = False; body.italic = False; body.font.color.rgb = INK; body.font.size = Pt(10.5); body.font.name = "Calibri"
    if note:
        tag = np.add_run("   [NEEDS YOUR INPUT — " + note + "]")
        tag.bold = True; tag.italic = False; tag.font.color.rgb = FLAG; tag.font.size = Pt(10.5); tag.font.name = "Calibri"
        tag.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return np


count = 0; flagged = 0
for idx, (ans, note) in A.items():
    add_answer(paras[idx], ans, note)
    count += 1
    if note:
        flagged += 1

d.save(DOC)
print(f"Inserted {count} answers ({flagged} flagged for your input).")
