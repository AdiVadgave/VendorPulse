"""Build VendorPulse_Development_Dependencies.docx.

Shell-facing development dependency document. Prose is written to read
naturally; the reference tables (versions, services, licences) are kept intact.

Usage: python _build_dependencies_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "VendorPulse_Development_Dependencies.docx"

SHELL_RED = RGBColor(0xC8, 0x10, 0x2E)
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x55, 0x55, 0x55)
HEAD_HEX = "C8102E"
ZEBRA_HEX = "F5F6F8"


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, *, bold=False, color=None, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade(hdr[i], HEAD_HEX)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
            if r_idx % 2 == 1:
                shade(cells[i], ZEBRA_HEX)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = SHELL_RED
    r.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = INK
    r.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p


def body(doc, text, *, muted=False, size=10):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    r.font.color.rgb = MUTED if muted else INK
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = "Calibri"
    r.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(2)
    return p


def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    # ---- Title block (single, clean) ----
    title = doc.add_paragraph()
    tr = title.add_run("VendorPulse — Development Dependency Document")
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.color.rgb = SHELL_RED
    tr.font.name = "Calibri"
    title.paragraph_format.space_after = Pt(2)

    meta = doc.add_paragraph()
    mr = meta.add_run("Prepared by Zensar for Shell   ·   Version 1.0   ·   10 June 2026   ·   Confidential")
    mr.font.size = Pt(9.5)
    mr.font.color.rgb = MUTED
    mr.font.name = "Calibri"
    meta.paragraph_format.space_after = Pt(10)

    # ---- 1. Introduction ----
    h1(doc, "1. Introduction")
    body(doc,
         "This document sets out what the team needs to build, deploy and run VendorPulse in "
         "Shell's environment: the developer hardware, the software stack with pinned versions, "
         "the Azure services the application depends on, and the commercial licences Shell will "
         "need to procure. It is a companion to the Dependencies & Access Requirements document, "
         "which covers identity and access provisioning and the associated lead times.")

    add_table(
        doc,
        ["", ""],
        [
            ["Scope", "Developer hardware, operating systems, runtimes, versioned libraries, the AI/agent stack, development and cloud tooling, and commercial licences."],
            ["Out of scope", "Identity and network provisioning steps and team allocation (see Dependencies & Access Requirements), and architecture rationale (see the Solution and Deployment Architecture documents)."],
            ["Audience", "Zensar engineering, and Shell IT Architecture, IT Security, Procurement and Cloud/IDT teams."],
        ],
        widths=[1.4, 5.6],
    )

    body(doc,
         "VendorPulse is a single-tenant web application. The backend is a FastAPI service "
         "(Python 3.11) running on Azure Container Apps; the frontend is a React 19 single-page "
         "application served from Azure Static Web Apps. It uses Entra ID for single sign-on and "
         "Microsoft Graph for Outlook, calendar and Teams, with PostgreSQL for storage and Key "
         "Vault for secrets. A Microsoft Agent Framework layer calls Microsoft Foundry to draft "
         "text, and every AI-generated output passes through a human-approval gate before any "
         "action is taken.")

    # ---- 2. Hardware ----
    h1(doc, "2. Hardware")
    h2(doc, "2.1 Developer workstation (per engineer)")
    body(doc,
         "There is no local GPU requirement — all model inference runs remotely in Foundry, so a "
         "standard development laptop is sufficient.")
    add_table(
        doc,
        ["Component", "Minimum", "Recommended"],
        [
            ["CPU", "Intel Core i5 11th Gen / AMD Ryzen 5 5000 (4c/8t)", "i7/i9 12th Gen+ / Ryzen 7–9 (8c+)"],
            ["RAM", "32 GB", "32–64 GB"],
            ["Storage", "512 GB SSD, ≥ 100 GB free", "1 TB NVMe"],
            ["GPU", "None (inference is remote to Foundry)", "—"],
            ["OS", "Windows 11 Pro / macOS 13+ / Ubuntu 22.04+", "+ WSL2 on Windows"],
            ["Network", "Broadband + Shell VPN", "Wired, low-latency"],
        ],
        widths=[1.1, 3.0, 2.9],
    )

    h2(doc, "2.2 Cloud runtime (Shell IDT-managed Azure)")
    body(doc,
         "VendorPulse runs entirely on Shell IDT-managed Azure; there is no on-premises footprint. "
         "The backend runs on Azure Container Apps and the frontend on Azure Static Web Apps, across "
         "three environments:")
    add_table(
        doc,
        ["Tier", "Backend (Container Apps)", "Frontend (Static Web Apps)", "PostgreSQL"],
        [
            ["Production", "Autoscale 2–N, zone-redundant", "Production", "General Purpose + HA"],
            ["Staging", "Autoscale 1–3", "Staging slot", "General Purpose (small)"],
            ["Development", "1 replica, scale-to-zero", "Preview", "Burstable"],
        ],
        widths=[1.0, 2.3, 2.0, 1.7],
    )

    # ---- 3. Software ----
    h1(doc, "3. Software (with versions)")
    body(doc,
         "Exact versions are pinned in requirements.txt and package-lock.json and scanned in CI, in "
         "line with Shell IRM control D2. Where a minimum is shown (≥), it is the project floor "
         "rather than a fixed pin.")

    h2(doc, "3.1 Runtimes and package managers")
    add_table(
        doc,
        ["Software", "Version"],
        [
            ["Python", "3.11.x"],
            ["Node.js", "20 LTS (≥ 20.19) or 22 LTS (≥ 22.12) — required by Vite 8"],
            ["npm / pip / Git", "10.x / 23+ / 2.40+"],
        ],
        widths=[2.0, 5.0],
    )

    h2(doc, "3.2 Backend — Python (requirements.txt)")
    add_table(
        doc,
        ["Package", "Version", "Purpose"],
        [
            ["fastapi", "0.115.6", "Web framework / REST API"],
            ["uvicorn[standard]", "0.32.1", "ASGI server"],
            ["pydantic / pydantic-settings", "2.10.4 / 2.7.0", "Validation / configuration"],
            ["httpx", "0.28.1", "Async HTTP (Microsoft Graph)"],
            ["openai", "≥ 1.50.0", "LLM SDK (Foundry Responses API)"],
            ["azure-ai-projects / azure-identity", "≥ 1.0.0 / ≥ 1.17.0", "Foundry client / Entra auth"],
            ["truststore", "≥ 0.10.0", "Corporate TLS trust"],
            ["python-dotenv", "1.0.1", "Local environment (dev only)"],
        ],
        widths=[2.4, 1.8, 2.8],
    )
    body(doc,
         "For the Shell production build we add the MAF SDK (pinned to an exact version), MSAL for "
         "app-only Entra authentication, the PostgreSQL drivers (psycopg/asyncpg) with SQLAlchemy "
         "and Alembic for migrations, and the Azure Key Vault and Monitor SDKs. The Google packages "
         "used in the proof of concept (Gmail and Forms) are dropped; their functions move to "
         "Microsoft Graph and a native form.")

    h2(doc, "3.3 Frontend — Node / React (package.json)")
    add_table(
        doc,
        ["Package", "Version", "Purpose"],
        [
            ["react / react-dom", "19.2.4", "UI"],
            ["vite", "8.0.1", "Build / dev server"],
            ["typescript", "~5.9.3", "Language"],
            ["react-router-dom", "7.13.2", "Routing"],
            ["zustand", "5.0.12", "State management"],
            ["recharts", "3.8.1", "Charts"],
            ["tailwindcss / @tailwindcss/vite", "4.2.2", "Styling"],
            ["date-fns / lucide-react / clsx / tailwind-merge", "4.1.0 / 1.7.0 / 2.1.1 / 3.5.0", "Utilities / icons"],
            ["eslint / typescript-eslint", "9.39.4 / 8.57.0", "Linting"],
        ],
        widths=[3.0, 2.2, 1.8],
    )

    h2(doc, "3.4 AI / agent stack")
    add_table(
        doc,
        ["Component", "Version / model"],
        [
            ["Microsoft Agent Framework (MAF) SDK", "1.x (≈1.8, June 2026) — pinned to an exact version"],
            ["Microsoft Foundry", "Responses API with content safety"],
            ["Model", "gpt-4.1 / gpt-4o family (the LLMProvider abstraction keeps Anthropic selectable)"],
            ["Authentication", "Entra ID (DefaultAzureCredential / app-only certificate) — no API keys in code"],
        ],
        widths=[3.0, 4.0],
    )

    h2(doc, "3.5 Developer and CI tooling")
    add_table(
        doc,
        ["Tool", "Version", "Purpose"],
        [
            ["VS Code (or JetBrains)", "latest", "IDE"],
            ["Claude Code (Anthropic agentic CLI)", "latest", "AI pair-programming assistant"],
            ["Azure CLI (az)", "2.60+", "Azure / Foundry / Graph authentication"],
            ["Docker", "24+", "Container image builds"],
            ["Bicep / Terraform", "latest", "Infrastructure as code"],
            ["ruff / eslint", "latest", "Linting in CI"],
            ["SonarQube/SonarCloud · Trivy · GitLeaks", "per Shell", "SAST · image/dependency scan · secrets scan"],
            ["GitHub Actions / Azure DevOps", "SaaS", "CI/CD (Shell Enterprise Agreement)"],
        ],
        widths=[3.0, 1.4, 2.6],
    )

    # ---- 4. Cloud platform ----
    h1(doc, "4. Cloud Platform and Managed Services")
    body(doc, "All services run in a single-tenant Azure deployment in West Europe.")
    add_table(
        doc,
        ["Service", "Tier", "Role"],
        [
            ["Azure Container Apps", "Autoscale", "FastAPI + MAF backend"],
            ["Azure Static Web Apps", "Standard", "React SPA (CDN)"],
            ["PostgreSQL Flexible Server", "GP + HA (prod)", "Datastore (BaseRepository seam)"],
            ["Key Vault", "Standard", "Secrets via Managed Identity"],
            ["Front Door + WAF", "Standard", "TLS, OWASP rules, origin lock"],
            ["App Insights + Azure Monitor", "Pay-as-ingested", "OpenTelemetry telemetry + audit mirror"],
            ["Container Registry (ACR)", "Standard", "Backend images"],
            ["Microsoft Foundry / Azure OpenAI", "Responses API", "LLM + content safety"],
            ["Microsoft Graph / Entra ID", "—", "Outlook/calendar/Teams · SSO, RBAC, identity"],
        ],
        widths=[2.6, 1.6, 2.8],
    )
    body(doc,
         "Private endpoints connect the backend to Key Vault, PostgreSQL and Foundry, and all "
         "infrastructure is defined as code (Bicep or Terraform).")

    # ---- 5. Licences ----
    h1(doc, "5. Licences and Subscriptions")
    body(doc,
         "Costs below are indicative; final figures depend on the Shell tier and usage. "
         "Open-source components carry no licence cost and require only a licence and security review.")

    h2(doc, "5.1 Commercial (procurement required)")
    add_table(
        doc,
        ["#", "Licence", "Tier", "Indicative cost", "Owner"],
        [
            ["LIC1", "Microsoft 365 (E3/E5) — service mailbox", "Per-user", "~$23 (E3) / ~$57 (E5) /user/mo", "Shell Messaging"],
            ["LIC2", "Microsoft Entra ID", "Included in M365", "—", "Shell Identity"],
            ["LIC3", "Azure subscription (Container Apps, SWA, PostgreSQL, Key Vault, Front Door, App Insights, ACR)", "Consumption", "~$450–$650/mo prod; +~$80–$120/mo non-prod", "Shell Cloud"],
            ["LIC4", "Microsoft Foundry / Azure OpenAI usage (gpt-4.1 / gpt-4o)", "Pay-per-token", "~$1,000/mo (≈$12k/yr) — Shell.AI + TRB approval", "Shell.AI"],
            ["LIC4-alt", "Anthropic Enterprise (LLM alternative via the abstraction)", "Enterprise", "Per contract + DPA", "Shell Procurement"],
            ["LIC5", "Azure DevOps / GitHub Enterprise", "SaaS", "Included in Microsoft EA", "Shell Engineering"],
            ["LIC6", "SonarQube / SonarCloud", "Per Shell", "Per licence", "Shell Engineering"],
            ["LIC7", "Claude Code (Anthropic)", "Pro / Max / Team seat", "~$20 (Pro) · $100–$200 (Max) /user/mo", "Zensar"],
            ["LIC8", "TLS certificate (public hostname)", "Shell PKI", "—", "Shell PKI"],
        ],
        widths=[0.7, 2.7, 1.1, 1.5, 1.0],
    )

    h2(doc, "5.2 Open-source (no licence cost — review only)")
    body(doc,
         "Python, FastAPI, Uvicorn, Pydantic and httpx; React, Vite, TypeScript, Zustand, Recharts "
         "and Tailwind; the Microsoft Agent Framework and Azure SDKs; and Trivy and GitLeaks. All "
         "are under MIT, BSD or Apache-2.0 licences. There are no copyleft (GPL) components in the "
         "current stack.")

    # ---- 6. Access & Identity ----
    h1(doc, "6. Access and Identity")
    body(doc,
         "This is a summary; the full breakdown, including provisioning steps, sits in the "
         "Dependencies & Access Requirements document.")
    add_table(
        doc,
        ["Ref", "Item", "Owner", "Lead time"],
        [
            ["I1–I2", "App registrations and certificate credentials", "Shell Entra ID", "1 week"],
            ["I3", "Admin consent: Mail.Send, Mail.Read, Calendars.ReadWrite, OnlineMeetings.ReadWrite.All, User.Read.All", "Shell Entra ID (global admin)", "2–4 weeks"],
            ["I5–I6", "Entra role groups and Application Access Policy (mailbox-scoped)", "Shell IAM / Exchange", "1–2 weeks"],
            ["A1–A5", "Azure subscription, resource groups, Container Apps, PostgreSQL, Key Vault", "Shell Cloud", "1–2 weeks"],
            ["Compliance", "AI Registry + ServiceNow + IRM IAQ; Shell.AI and TRB model approval", "Shell IRM", "Precedes production"],
        ],
        widths=[0.9, 3.4, 1.7, 1.0],
    )

    # ---- 7. Assumptions ----
    h1(doc, "7. Assumptions and Constraints")
    bullet(doc, "The target is the Shell production stack (MAF, Foundry, Azure and Outlook); the Google components from the proof of concept have been removed.")
    bullet(doc, "There is no local GPU — all inference is remote, in Foundry.")
    bullet(doc, "The LLM provider (Foundry or Anthropic) is a configuration choice through the LLMProvider abstraction.")
    bullet(doc, "In deployed environments, secrets come from Key Vault via Managed Identity; .env files are used for local development only.")
    bullet(doc, "The MAF SDK version is pinned and re-validated at build time.")

    foot = doc.add_paragraph()
    fr = foot.add_run("VendorPulse — Development Dependency Document · Zensar for Shell · v1.0 · 10 June 2026")
    fr.font.size = Pt(8)
    fr.font.color.rgb = MUTED
    fr.font.name = "Calibri"
    foot.paragraph_format.space_before = Pt(12)
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
