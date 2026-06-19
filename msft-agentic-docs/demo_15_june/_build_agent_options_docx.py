"""Build VendorPulse_Agent_Hosting_Options.docx — a short, human-voiced
recommendation note on how to run the agent layer on Microsoft Foundry.

Deliberately written as prose (an architect's memo), not the bulleted
house style, with a plain comparison table.

Usage: python _build_agent_options_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "VendorPulse_Agent_Hosting_Options.docx"

INK = RGBColor(0x20, 0x20, 0x20)
ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
MUTED = RGBColor(0x60, 0x60, 0x60)
HEAD_HEX = "1F3A5F"
ZEBRA_HEX = "F2F4F7"


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def cell_text(cell, text, *, bold=False, color=None, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def heading(doc, text, size=13, color=ACCENT, before=12, after=3):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


def para(doc, text, *, size=10.5, after=8, italic=False, color=INK):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    r.italic = italic
    r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    return p


def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    tr = title.add_run("Running the VendorPulse agent layer on Foundry: our recommendation")
    tr.bold = True
    tr.font.size = Pt(17)
    tr.font.color.rgb = ACCENT
    tr.font.name = "Calibri"
    title.paragraph_format.space_after = Pt(2)

    by = doc.add_paragraph()
    br = by.add_run("Zensar — Solution Architecture · June 2026")
    br.font.size = Pt(9.5)
    br.italic = True
    br.font.color.rgb = MUTED
    br.font.name = "Calibri"
    by.paragraph_format.space_after = Pt(10)

    para(doc,
         "All three of the options below run on the same Microsoft Foundry deployment: the same GPT-4o "
         "model, reached through the Responses API, inside Shell's own tenant. So this isn't really a "
         "decision about which model or which cloud. It comes down to two questions that matter more than "
         "they look at first glance: who orchestrates the agent, and who owns the prompt that goes to the "
         "model. For most projects those are details. For a governance tool they are the whole point, "
         "because the people who eventually sign this off will want to see exactly what the system does and "
         "be sure a human stays in control of anything that leaves the building.")

    para(doc,
         "Having weighed the three approaches Microsoft offers, our recommendation is to build on the Agent "
         "Framework SDK running inside our own application, and to keep Foundry's Hosted Agents in view as a "
         "follow-on once it reaches general availability. The portal-built agents are the one option we'd "
         "advise against, and the note below explains the reasoning.")

    heading(doc, "How the three options compare")

    headers = ["Consideration", "Portal agents", "Hosted agents", "In-process SDK (recommended)"]
    rows = [
        ["Where the agent code runs", "On Foundry", "On Foundry (managed compute)", "In our own backend"],
        ["Who controls the prompt", "Foundry assembles it", "We do", "We do"],
        ["Human-approval gate", "Has to be forced on top of an\n\"act-first\" design",
         "Kept exactly as designed", "Built into the architecture\n(the agent has no send tool)"],
        ["Auditability", "Prompt not shown to auditors\nverbatim", "Prompt lives in our repository",
         "Prompt lives in our repository"],
        ["Maturity", "Generally available", "Preview; GA expected early July",
         "Generally available"],
        ["Effect on the 4-week plan", "A rewrite; adds 2–3 weeks",
         "Roughly a week, plus a GA and\nregion dependency", "On plan"],
        ["Who runs the infrastructure", "Foundry", "Foundry", "We do (light for this workload)"],
    ]

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        cell_text(hdr[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=9.5)
        shade(hdr[i], HEAD_HEX)
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci, val in enumerate(row):
            cell_text(cells[ci], val, size=9, bold=(ci == 0))
            if ri % 2 == 1:
                shade(cells[ci], ZEBRA_HEX)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    heading(doc, "Building the agents in the Foundry portal")
    para(doc,
         "This is the option that usually looks most appealing at the start. There is nothing to host or "
         "operate, Microsoft maintains it, and it comes with a ready set of tool types. If we were building "
         "a general-purpose assistant we might well start here.")
    para(doc,
         "The trouble is that it works against the two things this particular product has to get right. "
         "Foundry assembles the instructions that go to the model, which means we cannot hand an auditor a "
         "single file and say \"this is the prompt, unchanged, every time\" — and that kind of transparency "
         "is exactly what IRM 3.492 and the EU AI Act expect of us. The Agent Service is also designed to "
         "act on its own: it wants to call tools and finish the job. Our whole safety model runs the other "
         "way round, with the agent only producing a draft and a person approving before anything is sent. "
         "We could rebuild that approval step on top of the portal, but we would be fighting the platform's "
         "natural behaviour and re-proving it to security with every release. On top of that, the logic "
         "would live in portal configuration rather than in our codebase, which makes it harder to review, "
         "harder to test, and harder to move later. Adopting it is effectively a rewrite of the agent layer, "
         "and it would push the timeline out by two to three weeks for a weaker compliance story.")

    heading(doc, "Hosting our own agents on Foundry")
    para(doc,
         "This is a genuinely good option, and it is the one to keep an eye on. Foundry runs the same agent "
         "code we would write anyway, so prompt ownership, the approval gate and the deterministic core all "
         "carry over untouched, while Foundry takes on the scaling, identity and tracing. It also supports a "
         "private network with no public egress, which sits well with Shell's controls.")
    para(doc,
         "The catch is timing. It is still in preview, with general availability expected in early July, "
         "and a preview feature is normally something Shell's risk and review boards will not let into "
         "production. The preview regions include Sweden Central but we have not confirmed West Europe, so "
         "there is a data-residency question to settle before committing. And moving to it is not a switch "
         "we flip — the agents would need to be packaged for the hosted runtime and given a route to their "
         "tools and data. None of that is hard, but it is a week of work that only makes sense once the "
         "feature is GA. Our view is to treat this as the natural next step rather than the starting point.")

    heading(doc, "Running the SDK inside our own application — what we recommend")
    para(doc,
         "This is where we would start. Because we write the agent code, the prompt sits in our repository "
         "where it can be versioned, diffed and shown to an auditor word for word. The approval gate stops "
         "being something we configure and becomes part of the design: the agent simply has no tool that "
         "sends anything, and the actual send happens in separate, deterministic code that only runs after a "
         "human clicks approve. The rest of the system — slot ranking, score validation, the workflow "
         "itself — stays as plain code, and the whole thing still works with the model switched off.")
    para(doc,
         "It is generally available today, it fits the four-week plan, and it is the easiest version for "
         "Shell to certify because every line is in our pipeline and nothing is hidden in a portal. It also "
         "keeps our options open: the same code lifts onto Hosted Agents when that is ready, and the "
         "provider abstraction means Anthropic stays selectable if procurement ever prefers it. The only "
         "real cost is that we operate the runtime ourselves, which for a steady internal workload is modest "
         "and which Hosted Agents removes for us later anyway.")

    heading(doc, "The two things we are not willing to give up")
    para(doc,
         "The first is the human in the loop. VendorPulse is built so that the AI drafts, a person approves, "
         "and only then does the system act. Keeping the agent unable to send anything on its own is not a "
         "feature we can compromise on, and it is the clearest reason to avoid the portal route, which is "
         "designed to act autonomously.")
    para(doc,
         "The second is being able to prove what the system did. Under IRM 3.492 and the EU AI Act, whoever "
         "reviews this will want to see the exact prompt and the exact path a decision took. With our own "
         "code that is straightforward — the prompt is in source control and the decisions are deterministic "
         "and logged. With agents assembled in the portal, that transparency is the thing we would lose.")

    heading(doc, "Where we land")
    para(doc,
         "Build on the in-process SDK now: it is generally available, it keeps us on the plan, it is fully "
         "auditable, and the approval gate is part of the architecture rather than bolted on. Move to "
         "Hosted Agents as a follow-on once it goes GA in July, since it is the same code and gives us "
         "Microsoft's managed compute without giving anything up. And steer away from the portal-built "
         "agents, which are the one path that would cost us both prompt ownership and the approval gate "
         "the product depends on.")

    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
