"""Update VendorPulse_PRD_Full.docx with the latest decisions, preserving format.

- Replaces PoC-era 'mock' / 'demo' integration phrasing with the real production
  integrations (Microsoft Graph, native in-app form).
- Adds a concise 'Technical Approach (MVP)' section reflecting the decisions made
  (two-VM hosting, App Gateway + WAF, single AI Service -> Foundry GPT-4o, MAF
  removed, Postgres, Entra SSO/RBAC, egress proxy, OTel).

Saves to a new file so the original is untouched.
"""
import copy
from docx import Document
from docx.oxml.ns import qn

SRC = "VendorPulse_PRD_Full.docx"
OUT = "VendorPulse_PRD_Full_updated.docx"

d = Document(SRC)


def set_text(i, text):
    p = d.paragraphs[i]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r._r.getparent().remove(r._r)
    else:
        p.add_run(text)


# template pPr (style 'List Paragraph' + numPr numId=2) copied from an existing bullet
TEMPLATE_PPR = d.paragraphs[11]._p.pPr


def add_bullet(text):
    p = d.add_paragraph()
    pe = p._p
    existing = pe.find(qn("w:pPr"))
    if existing is not None:
        pe.remove(existing)
    pe.insert(0, copy.deepcopy(TEMPLATE_PPR))
    p.add_run(text)
    return p


# ---- 1. Update PoC-era functional phrasing ----
set_text(8, "The MVP is delivered as a live, working product deployed in Shell's Azure environment.")
set_text(43, "Invite is shown in the UI for the organiser to approve; on approval it is sent via Outlook through Microsoft Graph.")
set_text(58, "On cycle trigger, the agent distributes scorecard forms to each assigned vendor and internal metric owner via Microsoft Graph email and a native in-app form.")
set_text(115, "Real-time or transcript-based note capture during the meeting; supports manual capture with AI structuring.")
set_text(121, "One-click distribution of minutes to attendees via Microsoft Graph (Outlook).")

# ---- 2. Append a concise Technical Approach section ----
d.add_paragraph("Technical Approach (MVP)", style="Heading 1")
d.add_paragraph(
    "The MVP is built and run entirely inside Shell's Azure environment, single-tenant and in the "
    "approved region. The business logic is deterministic and AI is limited to drafting text, with every "
    "AI output passing through a human-approval gate before anything is sent or actioned. The design "
    "follows Shell IRM 3.492 and EU AI Act expectations — human oversight, a deterministic core, "
    "in-tenant AI, and a full audit trail.",
    style="Normal")

d.add_paragraph("Architecture & Hosting", style="Heading 3")
add_bullet("Hosted in Shell's Azure subscription within a private network (VNet); the application has no public endpoint.")
add_bullet("Two Azure VMs: an application server (UI, authentication, the 12-state workflow engine, the approval gate) and a backend-services server (database, integrations, AI Service).")
add_bullet("Reached through an Application Gateway with a Web Application Firewall — TLS termination, OWASP rules, and origin-lock.")

d.add_paragraph("AI Approach", style="Heading 3")
add_bullet("A single in-house AI Service calls Azure AI Foundry (GPT-4o), which runs inside Shell's own tenant.")
add_bullet("The model only drafts text — narratives, summaries, minutes, suggested counters — and never computes figures or takes actions.")
add_bullet("Every AI output is a draft, held at the approval gate until a person approves it.")

d.add_paragraph("Data, Security & Identity", style="Heading 3")
add_bullet("Azure PostgreSQL is the system of record (cycles, scorecards, decisions, actions, audit); Blob Storage holds files such as transcripts and minutes; Key Vault holds all secrets, accessed via Managed Identity.")
add_bullet("Users sign in with Shell SSO (Entra ID, OIDC); access is role-based (Coordinator, Sponsor, Viewer) and enforced server-side.")
add_bullet("The application authenticates to Microsoft 365 with an app-only certificate.")

d.add_paragraph("Integrations & Observability", style="Heading 3")
add_bullet("Microsoft Graph handles Outlook email, calendar, and Teams meetings; scorecards are collected through a native in-app form.")
add_bullet("All outbound calls leave through Shell's controlled egress proxy.")
add_bullet("Telemetry and an immutable audit trail are captured via OpenTelemetry to Azure App Insights and Log Analytics.")

d.save(OUT)
print("Wrote", OUT)
