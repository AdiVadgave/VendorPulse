"""Build VendorPulse_Client_Meeting_Prep.docx — a plain-language briefing
that prepares the Zensar team for the prospective-client review.

It decodes the two shared docs (Solution Architecture + Infrastructure &
Software Requirements), tells the Foundry migration story, explains the
deployment models (on-prem vs cloud vs serverless), justifies the VM size,
gives a full glossary of every term, and arms the team with answers to the
questions a Product Owner / Solution Architect / AI Architect will ask.

Usage: python _build_client_prep_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "VendorPulse_Client_Meeting_Prep.docx"

SHELL_RED = RGBColor(0xC8, 0x10, 0x2E)
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x15, 0x80, 0x3D)
BLUE = RGBColor(0x1D, 0x4E, 0xD8)
HEAD_HEX = "C8102E"
ZEBRA_HEX = "F5F6F8"
BOX_HEX = "FDF3F4"


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, *, bold=False, color=None, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    # allow simple **bold** spans inside a cell
    parts = text.split("**")
    for idx, seg in enumerate(parts):
        if not seg:
            continue
        run = p.add_run(seg)
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        run.bold = bold or (idx % 2 == 1)
        if color is not None:
            run.font.color.rgb = color


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade(hdr[i], HEAD_HEX)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
            if r_idx % 2 == 1:
                shade(cells[i], ZEBRA_HEX)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = SHELL_RED
    r.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = INK
    r.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def body(doc, text, *, muted=False, size=10):
    p = doc.add_paragraph()
    parts = text.split("**")
    for idx, seg in enumerate(parts):
        if not seg:
            continue
        r = p.add_run(seg)
        r.font.size = Pt(size)
        r.font.name = "Calibri"
        r.bold = (idx % 2 == 1)
        r.font.color.rgb = MUTED if muted else INK
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    parts = text.split("**")
    for idx, seg in enumerate(parts):
        if not seg:
            continue
        r = p.add_run(seg)
        r.font.size = Pt(10)
        r.font.name = "Calibri"
        r.bold = (idx % 2 == 1)
        r.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(2)
    return p


def callout(doc, title, text):
    """A shaded one-cell 'remember this' box."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    shade(cell, BOX_HEX)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.name = "Calibri"
    r.font.color.rgb = SHELL_RED
    p2 = cell.add_paragraph()
    parts = text.split("**")
    for idx, seg in enumerate(parts):
        if not seg:
            continue
        rr = p2.add_run(seg)
        rr.font.size = Pt(10)
        rr.font.name = "Calibri"
        rr.bold = (idx % 2 == 1)
        rr.font.color.rgb = INK
    doc.add_paragraph()


def qa(doc, q, a):
    """Question (bold red) + answer."""
    p = doc.add_paragraph()
    rq = p.add_run("Q.  " + q)
    rq.bold = True
    rq.font.size = Pt(10)
    rq.font.name = "Calibri"
    rq.font.color.rgb = SHELL_RED
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(1)
    pa = doc.add_paragraph()
    parts = a.split("**")
    for idx, seg in enumerate(parts):
        if not seg:
            continue
        ra = pa.add_run(seg)
        ra.font.size = Pt(10)
        ra.font.name = "Calibri"
        ra.bold = (idx % 2 == 1)
        ra.font.color.rgb = INK
    pa.paragraph_format.space_after = Pt(4)
    pa.paragraph_format.left_indent = Inches(0.28)


def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    # ---- Title ----
    title = doc.add_paragraph()
    tr = title.add_run("VendorPulse — Client Meeting Prep Pack")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = SHELL_RED
    tr.font.name = "Calibri"
    title.paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph()
    sr = sub.add_run("Decoding the Solution Architecture & Infrastructure docs · the Foundry story · "
                     "deployment models · VM sizing · full glossary · anticipated Q&A")
    sr.font.size = Pt(11)
    sr.italic = True
    sr.font.color.rgb = MUTED
    sr.font.name = "Calibri"
    sub.paragraph_format.space_after = Pt(2)

    meta = doc.add_paragraph()
    mr = meta.add_run("Prepared for the Zensar team · for the prospective-client review (Product Owner · "
                      "Solution Architect · AI Architect) · Confidential")
    mr.font.size = Pt(9)
    mr.font.color.rgb = MUTED
    mr.font.name = "Calibri"
    meta.paragraph_format.space_after = Pt(8)

    # ---- 0. The 7 things to remember ----
    h1(doc, "0. Read this first — the 7 things to remember")
    body(doc, "If you remember nothing else walking into the room, remember these. Everything later in this "
              "pack just backs them up.")
    bullet(doc, "**What it is:** VendorPulse automates the **Quarterly Business Review (QBR)** process for vendor "
                "governance — a fixed **12-step workflow** that the software drives, with AI used only to **draft text**.")
    bullet(doc, "**Deterministic-first, AI-second:** every real decision (ranking slots, validating scores, moving "
                "the workflow forward) is **plain code**. The AI never decides or acts — it only writes drafts a "
                "human approves. This is the single most important sentence for an AI-governance client.")
    bullet(doc, "**Human-approval gate (HITL):** nothing leaves the system — no email, no invite — until a person "
                "clicks Approve. The AI cannot send anything by itself.")
    bullet(doc, "**Where the AI runs:** **Azure AI Foundry, inside Shell's own tenant** (the GPT-4o model), reached "
                "over a **private** connection. No data goes to the public internet or trains any model.")
    bullet(doc, "**How it's hosted:** **100% cloud, in Shell's own Azure subscription. Not on-premises. Not "
                "serverless.** The current design runs the app on a single **Azure virtual machine (D4as v6 / "
                "D8as v6)** inside a private network.")
    bullet(doc, "**What changed recently:** we started by calling the model directly from FastAPI; on your steer we "
                "moved to the **Azure AI Foundry + Microsoft Agent Framework (MAF)** approach. A working **proof of "
                "concept is done**.")
    bullet(doc, "**No GPU anywhere:** all the heavy AI compute happens remotely in Foundry, so neither the server "
                "nor the developer laptops need a graphics card.")

    callout(doc, "ONE-LINE ANSWER if asked 'is this serverless / on-prem / cloud?'",
            "\"It runs entirely in Shell's own Azure cloud — single-tenant, inside your region and your network. "
            "It is **not** on-premises (no hardware in a Shell building) and the current design is **not** "
            "serverless: it's a right-sized **Azure virtual machine** we fully control, which keeps the security "
            "review simple. We can move it to a managed/serverless Azure service later without changing the code.\"")

    # ---- 1. The project in plain English ----
    h1(doc, "1. The project in plain English")
    body(doc, "VendorPulse is a **vendor-governance workflow app**. Large organisations run **QBRs** — Quarterly "
              "Business Reviews — with their key vendors: a structured meeting where both sides review how the "
              "vendor performed (scorecards), agree actions, and prepare for the next quarter. Coordinating one is "
              "a lot of manual chasing: find a meeting slot everyone can make, send invites, collect scorecards, "
              "compile them, brief leadership, prep for the meeting, capture minutes. VendorPulse turns that into a "
              "guided, auditable assembly line.")
    body(doc, "The work moves through a **12-state workflow** (created → attendees refreshed → availability "
              "collected → meeting scheduled → scorecards requested → collected → compiled → internal alignment → "
              "vendor prep → meeting → post-meeting → archived). The workflow can only move **forward, one step at "
              "a time** — it can never skip or go backwards. That rigidity is a feature: it makes the process "
              "predictable and auditable, which is exactly what a governance tool needs.")
    body(doc, "There are **6 functional modules (A–F)** — Scheduling, Scorecard, Alignment, Vendor Prep, Meeting, "
              "Analytics. Each has an **AI agent** that drafts the human-readable text for that step (an invite, a "
              "brief, meeting minutes, a leadership summary). The agent **never** does the maths or makes the "
              "decision — that's deterministic code. It only writes the words, and a human approves them.")

    callout(doc, "The mental model to carry into the room",
            "Think of VendorPulse as a **factory conveyor belt with a smart assistant standing next to it**. The "
            "belt (the workflow) is rigid and deterministic. The assistant (the AI) can draft a nicely-worded "
            "email or summary when asked, but it can't move the belt, change a number, or post anything — a human "
            "supervisor signs off every output. Governance clients buy the conveyor belt; the AI is a convenience "
            "on top, fully fenced in.")

    # ---- 2. Decoding the two shared documents ----
    h1(doc, "2. The two shared documents, decoded")

    h2(doc, "2.1 Solution Architecture — the 5 numbered layers")
    body(doc, "The architecture diagram is read **left to right in 5 zones**. Here is what each one means in plain "
              "language:")
    add_table(
        doc,
        ["#", "Zone", "What it actually means"],
        [
            ["1", "Client", "The user's **web browser**. Three kinds of user — VMO Coordinator, Sponsor, Viewer — "
                            "all log in with their normal Shell company account (**Entra SSO**). Nothing is installed "
                            "on their machine."],
            ["2", "Edge", "**Azure Front Door + WAF** — the secure front gate. It terminates HTTPS, applies a "
                          "**Web Application Firewall** (blocks common web attacks), and only lets traffic through to "
                          "our server (\"origin-lock\")."],
            ["3", "Shell Azure Subscription", "Shell's own slice of Azure, **inside the approved region**. Here lives a "
                              "private network (VNet) containing the **Azure VM** that runs the whole app: the React "
                              "screen, the login/permissions, the 12-state WorkflowEngine, the **approval gate**, the "
                              "GraphService, and the **MAF agent layer** that talks to Foundry."],
            ["4", "Data Tier", "The app's private back room, reachable only over **Private Link** (no public access): "
                               "**PostgreSQL** (the database), **Key Vault** (passwords & certificates), **Blob "
                               "Storage** (meeting minutes/transcripts), **App Insights + Log Analytics** (the audit "
                               "trail), and **Azure AI Foundry** (GPT-4o) — all inside Shell's tenant."],
            ["5", "External", "The only things reached outside, over outbound HTTPS via Shell's egress proxy: "
                              "**Microsoft Graph** (to send Outlook mail, manage calendars, create Teams meetings) and "
                              "**Entra ID** (to verify who the user is)."],
        ],
        widths=[0.3, 1.5, 5.2],
    )
    callout(doc, "Watch-out (be honest if asked)",
            "The diagram shows an **Azure App Service** icon at the top of zone 3 but labels the compute box "
            "**Azure VM**, and the Infrastructure doc specifies a VM (D4as v6 / D8as v6). If a sharp architect "
            "spots this, the honest answer is: **\"the compute target is an Azure VM; the App Service icon is a "
            "leftover from an earlier draft. Both are valid Azure hosting options and the application code is "
            "identical either way.\"** See §4 for why VM was chosen.")

    h2(doc, "2.2 Infrastructure & Software Requirements — what it lists")
    body(doc, "This is the shopping list: the cloud services, the software versions, the licences, and the access "
              "Shell must provision. The headline entry — and the one leadership keeps asking about — is the very "
              "first row: the backend runs on an **Azure VM, size D4as v6 or D8as v6** (covered in detail in §4 "
              "and §5). Everything else (Static Web Apps, PostgreSQL, Key Vault, Front Door, Foundry, Graph) "
              "matches the architecture diagram. The software section pins exact versions (Python 3.11, FastAPI "
              "0.115.6, React 19, the MAF SDK ≈1.8, etc.) so the build is reproducible and passes Shell's security "
              "scanning.")

    # ---- 3. The Foundry story ----
    h1(doc, "3. The Foundry story — where we came from, where we are")
    body(doc, "You will be asked why the design mentions Foundry, MAF, and an abstraction layer. Here is the "
              "honest, simple narrative.")
    add_table(
        doc,
        ["Stage", "What we did", "Why it matters to the client"],
        [
            ["**Start (PoC v1)**", "A FastAPI backend that **called the model directly** — raw API calls to the "
                       "Foundry / Azure OpenAI **Responses API**, with our own hand-written tool-calling loop.",
                       "Proved the idea worked, but every control (approval, safety, logging) was hand-rolled by us "
                       "— more code for Shell's security team to review."],
            ["**Your steer (last meeting)**", "Use **Azure AI Foundry** properly — i.e. build on the **Microsoft "
                       "Agent Framework (MAF)** SDK running over the Foundry Responses API, instead of raw calls.",
                       "Foundry/MAF provide **platform-built** human-in-the-loop approval, content safety filters, "
                       "and tracing — fewer hand-rolled controls, which is easier to get through Shell's review."],
            ["**Now (PoC v2 — done)**", "A working proof of concept on the **MAF + Foundry Responses API** across "
                       "2 agents, with the approval gate proven (GitHub issue #13). You may see a branch named "
                       "**poc/scheduling-foundry-responses**.",
                       "De-risks the build: the agentic pattern, tool-calling, and approval gate are all proven on "
                       "real Foundry before any production spend."],
            ["**Decision**", "**Production builds on the MAF SDK**, with the direct-Responses path kept as a "
                       "fallback. An **LLMProvider abstraction** keeps Anthropic/Claude selectable if procurement "
                       "prefers it.",
                       "We're not locked to one vendor; the model choice is a configuration switch, not a rewrite."],
        ],
        widths=[1.3, 3.0, 2.7],
    )
    h2(doc, "The three terms people confuse — say them like this")
    bullet(doc, "**Azure OpenAI** = just the **model behind an API** (you send text, you get text). The plumbing.")
    bullet(doc, "**Azure AI Foundry** = the **whole platform** around those models — model catalog, agents, content "
                "safety, tracing, deployments. Foundry is where our GPT-4o deployment lives, inside Shell's tenant. "
                "We reach it through its **Responses API**.")
    bullet(doc, "**Microsoft Agent Framework (MAF)** = the **open-source SDK** we write our agents with "
                "(`agent_framework.Agent` + `@tool` functions). MAF is our code; it calls Foundry. It gives us the "
                "agent loop, tool-calling, and a built-in approval mode for free.")
    callout(doc, "If asked 'so are you using Azure OpenAI or Azure AI Foundry?'",
            "\"Both names describe the same Microsoft stack at different layers. Our agent code is written with the "
            "**Microsoft Agent Framework**; it calls the **Azure AI Foundry Responses API**; and the model serving "
            "behind that is the **Azure OpenAI GPT-4o** family — all hosted inside Shell's own Foundry tenant.\"")

    # ---- 4. Deployment models explained ----
    h1(doc, "4. Deployment models, explained for leadership")
    body(doc, "This is the section to clear up the on-prem / cloud / serverless confusion. There is a **spectrum** "
              "of how much you manage yourself versus how much the cloud manages for you. From most-you to most-cloud:")
    add_table(
        doc,
        ["Model", "Plain-English analogy", "Who manages what", "Is VendorPulse this?"],
        [
            ["**On-premises (on-prem)**", "You **own the house** — buy the land, build it, fix the boiler.",
             "**You** own the physical servers in your own building/datacenter: hardware, power, cooling, patching.",
             "**No.** Nothing runs in a Shell building. Zero physical hardware for us."],
            ["**IaaS — cloud VM** *(our current design)*", "You **rent an empty apartment** — the building is "
             "managed, you furnish and clean inside.",
             "Azure owns the hardware; **you** manage the virtual machine's OS, patching, and your app on top.",
             "**Yes — this is it today.** An **Azure VM (D4as v6/D8as v6)** in Shell's Azure."],
            ["**PaaS — managed platform**", "You **rent a serviced apartment** — furniture and cleaning included, "
             "you just live there.",
             "Azure manages the OS and runtime; you just deploy your code/container. e.g. Azure App Service, "
             "Container Apps.",
             "**Not currently**, but a valid future move — the code doesn't change. (The earlier draft used this.)"],
            ["**Serverless / FaaS**", "You **stay in a hotel** — pay per night, walk away, staff handle everything.",
             "Azure manages everything incl. scaling; you pay per execution and it can **scale to zero**. e.g. "
             "Azure Functions.",
             "**No.** The current design is a constantly-running VM, not pay-per-call serverless."],
        ],
        widths=[1.5, 1.9, 2.0, 1.6],
    )
    h2(doc, "So, in one breath:")
    body(doc, "**VendorPulse is cloud-hosted IaaS, single-tenant, inside Shell's own Azure subscription and region. "
              "It is not on-premises and it is not serverless.** It's a virtual machine we control, sitting in a "
              "private network, talking to managed Azure services (database, secrets, AI) over private links.")
    h2(doc, "Why a plain VM and not serverless / a managed platform?")
    bullet(doc, "**Simplest security review (the big one for Shell):** a single VM in a VNet is the easiest shape "
                "to reason about, lock down, and get past IRM / IT-Security. Fewer moving managed services = fewer "
                "things to assess.")
    bullet(doc, "**Full control & no preview dependencies:** we own the OS and runtime; nothing depends on a "
                "preview-tier Azure feature.")
    bullet(doc, "**Predictable cost & behaviour:** an always-on VM has a flat, predictable monthly cost and no "
                "cold-start latency — fine for an internal governance tool with steady, modest traffic.")
    bullet(doc, "**The code is portable:** because the app is a standard container/FastAPI app, moving it later to "
                "Container Apps (PaaS) or a Foundry Hosted Agent (managed) is a **hosting swap, not a rewrite**. We "
                "keep that option open.")
    callout(doc, "If they push: 'serverless would be cheaper / auto-scale — why not?'",
            "\"For a steady internal workload with a human approval step in the loop, the elastic scaling and "
            "scale-to-zero of serverless buy us little, while a single VM is far simpler for your security team to "
            "certify. The application is containerised, so if usage grows we can move to Azure Container Apps or a "
            "Foundry Hosted Agent — same code — and get autoscale then. We chose the option that ships fastest "
            "through Shell governance.\"")

    # ---- 5. VM sizing ----
    h1(doc, "5. VM sizing — exactly what to say")
    body(doc, "The Infrastructure doc specifies an **Azure D4as v6** or **D8as v6** virtual machine. Here is what "
              "those names mean and how to defend the choice.")
    h2(doc, "5.1 Decoding the VM name")
    add_table(
        doc,
        ["Part of the name", "Meaning"],
        [
            ["**D**", "**D-series = general-purpose** VM (balanced CPU-to-memory). The right family for a normal web "
                      "/ API workload — not compute-heavy, not memory-heavy."],
            ["**4 / 8**", "Number of **vCPUs** (virtual CPU cores): **4** or **8**."],
            ["**a**", "Runs on **AMD** (EPYC) processors — strong price/performance."],
            ["**s**", "**Premium SSD capable** — fast, reliable disk."],
            ["**v6**", "**Generation 6** — the current, efficient hardware generation."],
        ],
        widths=[1.6, 5.4],
    )
    h2(doc, "5.2 The actual specs")
    add_table(
        doc,
        ["Size", "vCPUs", "RAM", "Good for", "Indicative cost*"],
        [
            ["**D4as v6**", "4 cores", "16 GB", "Dev / Staging, and Production at expected QBR volumes (steady, "
                            "modest concurrent users).", "~$140–200 / month"],
            ["**D8as v6**", "8 cores", "32 GB", "Production headroom if many coordinators work concurrently, or "
                            "for comfortable margin.", "~$280–400 / month"],
        ],
        widths=[1.1, 0.8, 0.8, 3.0, 1.3],
    )
    body(doc, "*Cost is pay-as-you-go indicative for West Europe and drops materially with a 1- or 3-year "
              "reserved-instance commitment — present figures as \"order of magnitude,\" not a quote.", muted=True)
    h2(doc, "5.3 Why this size is right — the reasoning to give")
    bullet(doc, "**The heavy lifting isn't here.** The AI inference runs remotely in **Foundry**, and the database "
                "runs as a separate managed service. The VM only runs a lightweight **FastAPI** web app and the MAF "
                "orchestration code — that is not CPU- or memory-hungry.")
    bullet(doc, "**No GPU needed** — and we should say this proactively, because people assume \"AI = expensive "
                "GPU server.\" There is **no GPU** on the VM or on developer laptops; all model compute is Foundry's.")
    bullet(doc, "**4 vCPU / 16 GB comfortably handles** a FastAPI app serving an internal governance team. The "
                "workload is bursty-but-small (a coordinator drafting an invite, approving a brief) — not "
                "high-throughput.")
    bullet(doc, "**Vertical headroom is trivial:** if it ever needs more, resizing a VM to the next size (e.g. "
                "D4→D8) is a few-minute reboot in the portal — no re-architecting.")
    bullet(doc, "**Production gets resilience** via Azure (zone-redundancy / availability), and the database is "
                "**High-Availability** — so the small VM size doesn't compromise uptime.")
    callout(doc, "If asked 'what spec / how big a server?' — the crisp answer",
            "\"A general-purpose Azure VM — **4 vCPUs and 16 GB RAM (D4as v6)** for most environments, scaling to "
            "**8 vCPU / 32 GB (D8as v6)** in production if we want headroom. No GPU is needed because all AI "
            "inference happens in Foundry, and the database is a separate managed service. The VM just runs the "
            "FastAPI app — it's a light workload, and we can resize in minutes if volumes grow.\"")

    # ---- 6. Full glossary ----
    h1(doc, "6. The full glossary — every term, in plain language")
    body(doc, "Grouped so you can find a term fast. For each: what it is, and (where useful) why it's in our design.")

    h2(doc, "6.1 Deployment & hosting")
    add_table(
        doc,
        ["Term", "Plain meaning"],
        [
            ["On-premises / on-prem", "Physical servers you own and run in your own building/datacenter. **Not used "
                                      "here.**"],
            ["Cloud / public cloud", "Computing rented from a provider (here, Microsoft **Azure**) over the "
                                     "internet — no hardware to own."],
            ["IaaS (Infrastructure as a Service)", "You rent a **virtual machine**; the cloud runs the hardware, you "
                                                   "manage the OS + app. **VendorPulse's current model.**"],
            ["PaaS (Platform as a Service)", "The cloud manages the OS + runtime; you just deploy code/containers "
                                             "(e.g. App Service, Container Apps)."],
            ["SaaS (Software as a Service)", "Finished software you just log into (e.g. Microsoft 365). VendorPulse "
                                             "would be delivered to Shell as an internal app, not sold as SaaS."],
            ["Serverless / FaaS", "Pay-per-execution compute that scales to zero (e.g. Azure Functions). The cloud "
                                  "handles all scaling. **Not the current design.**"],
            ["Single-tenant", "The whole deployment is **dedicated to Shell** — Shell's data, tenant, region, no "
                              "sharing with other customers."],
            ["Multi-tenant", "One shared deployment serving many customers (opposite of above). Not how this is built."],
            ["Virtual Machine (VM)", "A computer that exists as software inside Azure's datacenter — you get an OS "
                                     "and full control, without owning hardware."],
            ["vCPU", "A **virtual CPU core** — the unit of processing power assigned to a VM."],
            ["Autoscale / replicas", "Automatically adding/removing copies of the app as load changes. A managed "
                                     "/serverless feature; a single VM scales by resizing instead."],
            ["High Availability (HA)", "Running redundantly so a single failure doesn't cause downtime. Our "
                                       "**PostgreSQL** is HA in production."],
            ["Zone-redundant", "Spread across separate Azure datacenters in a region so one datacenter outage "
                               "doesn't take you down."],
            ["Region / data residency", "The geographic location of the datacenter (e.g. **West Europe**). Keeping "
                                        "data in an approved region is a Shell compliance requirement."],
            ["IaC (Infrastructure as Code)", "Defining all the cloud setup in text files (**Bicep / Terraform**) so "
                                             "environments are reproducible and reviewable."],
        ],
        widths=[2.1, 4.9],
    )

    h2(doc, "6.2 Azure services in the docs")
    add_table(
        doc,
        ["Term", "Plain meaning"],
        [
            ["Azure subscription / resource group", "Shell's billing account in Azure / a labelled folder grouping "
                                                    "related resources."],
            ["Azure VM", "The virtual machine running the FastAPI + MAF backend (D4as v6 / D8as v6)."],
            ["Azure Static Web Apps", "A cheap, fast Azure service for hosting the built **React** front-end, "
                                      "delivered globally via CDN."],
            ["Azure Container Apps / App Service", "Managed (PaaS) ways to run the backend without managing a VM — "
                                                   "**alternatives** to the VM, same code."],
            ["Azure Container Registry (ACR)", "A private store for the app's **Docker container images**."],
            ["Azure PostgreSQL Flexible Server", "The managed **database** (open-source PostgreSQL). \"Flexible "
                                                 "Server\" is the deployment option; **GP** = General Purpose tier, "
                                                 "**Burstable** = cheap dev tier."],
            ["Azure Key Vault", "A secure safe for **secrets** — passwords, API keys, certificates — so they're "
                                "never in code or config files."],
            ["Azure Front Door + WAF", "The global secure entry point. **WAF** (Web Application Firewall) blocks "
                                       "common web attacks using **OWASP** rules."],
            ["OWASP", "An industry-standard list of the top web-application security risks; the WAF enforces rules "
                      "against them."],
            ["App Insights / Azure Monitor / Log Analytics", "Azure's **observability** stack — collects logs, "
                                                             "metrics, and traces; also our immutable **audit** mirror."],
            ["Azure Blob Storage", "Cheap object storage for files (future: meeting minutes / transcripts)."],
            ["VNet (Virtual Network)", "A **private network** in Azure; our VM lives inside it, isolated from the "
                                       "public internet."],
            ["Private Endpoint / Private Link", "A private, internal-only connection to an Azure service (database, "
                                                "Key Vault, Foundry) so traffic **never traverses the public "
                                                "internet**."],
            ["Egress proxy", "Shell's controlled outbound gateway — the only way the app reaches external services "
                             "(Graph, Entra)."],
            ["Origin-lock / Service Tag", "Configuration ensuring the backend only accepts traffic from Front Door, "
                                          "not directly from the internet."],
            ["TLS / TLS 1.2+", "The encryption protocol behind HTTPS that protects data in transit."],
        ],
        widths=[2.4, 4.6],
    )

    h2(doc, "6.3 Identity & access")
    add_table(
        doc,
        ["Term", "Plain meaning"],
        [
            ["Microsoft Entra ID", "Microsoft's identity service — **formerly called Azure Active Directory (Azure "
                                   "AD)**. It's how users log in and how the app proves its identity."],
            ["SSO (Single Sign-On)", "Users log in with their existing Shell account — no separate password for "
                                     "VendorPulse."],
            ["OIDC (OpenID Connect)", "The standard protocol behind that SSO login."],
            ["RBAC (Role-Based Access Control)", "Permissions by role — e.g. Coordinator vs Viewer — so people only "
                                                 "see/do what their role allows."],
            ["JWT (JSON Web Token)", "The signed digital \"badge\" a user carries after login to prove who they are "
                                     "on each request."],
            ["Managed Identity", "An automatic, password-less identity Azure gives the VM so it can access Key Vault, "
                                 "the database, and Foundry **without any stored secret**."],
            ["MSAL", "Microsoft's auth library used for the app-to-Graph login."],
            ["App registration", "The app's identity record in Entra, defining what it's allowed to do."],
            ["Client-credentials flow", "App-to-service login (no human involved) — lets VendorPulse send mail or "
                                        "read calendars unattended."],
            ["Admin consent", "A Shell global admin's one-time approval of the permissions the app requests (a "
                              "**2–4 week** lead-time item)."],
            ["Application Access Policy", "An Exchange control limiting the app's mail permission to **one specific "
                                          "mailbox**, not everyone's."],
            ["App-only / certificate auth", "The production app authenticates as itself using a **certificate** "
                                            "(stronger than a secret), replacing the PoC's pasted 1-hour token."],
        ],
        widths=[2.2, 4.8],
    )

    h2(doc, "6.4 Microsoft Graph & Microsoft 365")
    add_table(
        doc,
        ["Term", "Plain meaning"],
        [
            ["Microsoft Graph", "The single API to reach Microsoft 365 data — **Outlook mail, calendars, Teams "
                                "meetings, user lookups**. VendorPulse uses it to send invites and schedule meetings."],
            ["Mail.Send / Mail.Read", "Graph permissions: send email / read replies."],
            ["Calendars.ReadWrite", "Graph permission: read free/busy and create calendar events."],
            ["OnlineMeetings.ReadWrite.All", "Graph permission: create Teams meetings."],
            ["User.Read.All", "Graph permission: look up users in the directory (for attendee lists)."],
            ["Microsoft 365 E3 / E5", "Microsoft's productivity+security licence tiers; needed for the service "
                                      "mailbox the app sends from. E5 adds advanced security/compliance."],
        ],
        widths=[2.2, 4.8],
    )

    h2(doc, "6.5 AI / agent stack")
    add_table(
        doc,
        ["Term", "Plain meaning"],
        [
            ["LLM (Large Language Model)", "The AI that generates text (here, **GPT-4o**). In VendorPulse it only "
                                           "**drafts** wording — it never decides or acts."],
            ["GPT-4o / gpt-4.1", "Specific OpenAI model versions, deployed inside Shell's Foundry. **GA** "
                                 "(Generally Available) = stable, production-grade, not preview."],
            ["Azure OpenAI", "The Azure service that serves OpenAI models behind an API — the model layer."],
            ["Azure AI Foundry", "Microsoft's **platform** around those models: model catalog, agents, content "
                                 "safety, tracing. Our GPT-4o lives here, in Shell's tenant."],
            ["Responses API", "The Foundry/OpenAI API our agents call to get model output and tool-calls. The single "
                              "entry point we use."],
            ["Microsoft Agent Framework (MAF)", "The open-source **SDK** we build our agents with. Provides the "
                                                "agent loop, tool-calling, and a built-in approval mode. **Our code**, "
                                                "calling Foundry."],
            ["Agent", "A small AI-driven routine for one module (e.g. SchedulingAgent) that, given a prompt, can call "
                      "a set of tools and return a structured result."],
            ["Tool / tool-calling / @tool", "Functions the agent is allowed to call (e.g. \"rank slots\", \"draft "
                                            "invite\"). The model picks which to call; our code runs them. The result "
                                            "is **structured data**, not free text."],
            ["HITL (Human-in-the-loop) / approval_mode", "The control where a human must approve before a sensitive "
                                                         "action runs. VendorPulse gates this both in the app **and** "
                                                         "natively in MAF (belt-and-suspenders)."],
            ["Content safety / content filters", "Foundry's built-in filters that block harmful or unsafe model "
                                                 "input/output. **XPIA** = protection against prompt-injection attacks."],
            ["OpenTelemetry (OTel) / tracing", "An open standard for emitting logs/traces. MAF emits these by "
                                               "default; they flow to App Insights for the audit trail."],
            ["Hallucination", "When an LLM confidently states something false. We mitigate it by keeping all facts/"
                              "numbers/IDs in **deterministic code**, not the model."],
            ["Deterministic", "Code that always gives the same output for the same input (the opposite of an LLM's "
                              "variability). All VendorPulse decisions are deterministic."],
            ["LLMProvider abstraction", "A config switch that lets us swap the model vendor (Foundry **or** "
                                        "Anthropic/Claude) without changing application code."],
            ["AgentResponse / adapter", "The single, fixed response shape every agent returns, so the front-end never "
                                        "has to guess the output format."],
            ["Foundry Hosted Agents", "A (preview) Foundry option to host **our** agent code as a managed endpoint — "
                                      "a possible future upgrade from the VM, same code."],
            ["Claude / Anthropic", "An alternative LLM vendor, selectable via the abstraction. **Claude Code** is "
                                   "separately the AI coding assistant the dev team uses to build faster."],
        ],
        widths=[2.3, 4.7],
    )

    h2(doc, "6.6 Application & developer stack")
    add_table(
        doc,
        ["Term", "Plain meaning"],
        [
            ["FastAPI", "The Python web framework the backend is built on — serves the REST API."],
            ["Python 3.11", "The backend programming language/version."],
            ["Uvicorn / ASGI", "The server that runs the FastAPI app."],
            ["Pydantic", "Library that validates data shapes — guarantees inputs/outputs match the defined schema."],
            ["React 19 / SPA", "The front-end framework. **SPA** = Single-Page Application: the UI runs in the "
                               "browser and updates without full page reloads."],
            ["Vite / TypeScript / Zustand / Recharts / Tailwind", "Front-end build tool / typed JavaScript / state "
                                                                  "management / charts / styling."],
            ["WorkflowEngine / 12-state machine", "The deterministic core that enforces the **forward-only** 12-step "
                                                  "process — the single source of truth for what's allowed next."],
            ["Module A–F / the 6 agents", "Scheduling, Scorecard, Alignment, Vendor Prep, Meeting, Analytics — each a "
                                          "functional area with its own agent."],
            ["Repository pattern / BaseRepository", "The single code layer that touches storage. Swapping JSON files "
                                                    "→ PostgreSQL means changing **only** this layer."],
            ["PostgreSQL / SQLite / JSON files", "The production database / the lightweight dev database / the "
                                                 "original file-based storage being replaced."],
            ["SQLAlchemy / Alembic", "Database toolkit / database-migration tool for evolving the schema safely."],
            ["Magic-link / one-time token", "A secure single-use link emailed to a stakeholder to open the scorecard "
                                            "form without a full account."],
            ["CI/CD", "Continuous Integration / Continuous Delivery — the automated pipeline that lints, tests, "
                      "builds, scans, and deploys code (GitHub Actions / Azure DevOps)."],
            ["Docker / container / image", "Packaging the app with everything it needs to run identically anywhere; "
                                           "the **image** is the packaged artifact."],
            ["SAST / SonarQube / Trivy / GitLeaks", "Security scanning in CI: static code analysis / image & "
                                                    "dependency scanning / secret-leak detection."],
            ["ruff / eslint / tsc", "Automated code-quality/linting checks for Python and TypeScript."],
        ],
        widths=[2.4, 4.6],
    )

    h2(doc, "6.7 Governance & compliance (Shell-specific)")
    add_table(
        doc,
        ["Term", "Plain meaning"],
        [
            ["QBR (Quarterly Business Review)", "The recurring vendor-performance meeting VendorPulse orchestrates."],
            ["VMO (Vendor Management Office)", "The team that runs vendor governance — the primary users."],
            ["IRM 3.492", "Shell's internal control standard for AI systems — the rulebook this design maps to."],
            ["NIST AI RMF / ISO 42001", "External AI risk-management / AI-management-system standards that IRM 3.492 "
                                        "is built on."],
            ["EU AI Act", "European regulation on AI use; requires transparency (e.g. labelling AI-generated content)."],
            ["AI Registry / ServiceNow / IAQ", "Shell's mandatory AI-system registration and risk-assessment "
                                               "(IRM Assessment Questionnaire) steps — must be done **before "
                                               "production**."],
            ["Shell.AI / TRB", "Shell's AI governance body / Technology Review Board — must approve the model and "
                               "design before production."],
            ["IDT", "Shell's Information & Digital Technology org — production **must** run on IDT-managed Azure/"
                     "Foundry."],
            ["DPA / no-training assurance", "Data Processing Agreement; the contractual guarantee that Shell's data "
                                            "is **not used to train** any model."],
            ["Data classification", "Labelling data by sensitivity — scorecards are likely **Commercially "
                                    "Sensitive**, attendee data is **PII/GDPR**. (No SOX / export-controlled data is "
                                    "in scope.)"],
        ],
        widths=[2.3, 4.7],
    )

    # ---- 7. Anticipated Q&A ----
    h1(doc, "7. Anticipated questions & crisp answers")
    body(doc, "Organised by who's likely to ask. Answers are written to be said out loud.")

    h2(doc, "7.1 From the Product Owner (value, scope, users)")
    qa(doc, "What does this actually save us?",
       "It removes the manual chasing in every QBR — scheduling, scorecard collection, compiling, briefing, "
       "minutes — into one guided, auditable workflow. Coordinators approve drafts instead of writing from scratch; "
       "leadership gets consistent briefs; nothing falls through the cracks because the workflow enforces every step.")
    qa(doc, "Who are the users and how do they log in?",
       "Three roles — VMO Coordinator, Sponsor, and Viewer — all logging in with their normal Shell account via "
       "single sign-on. No new passwords, permissions are role-based.")
    qa(doc, "How much of this is AI 'making decisions'?",
       "**None.** Every decision — slot ranking, score validation, workflow progression — is deterministic code. "
       "The AI only drafts human-readable text, and a person approves it before anything happens.")
    qa(doc, "What's live today vs still to build?",
       "A working proof of concept proves the agentic pattern and the approval gate on real Foundry (2 agents). "
       "Production work — full Azure infra, identity hardening, database migration, porting all agents, and "
       "clearing Shell's compliance gates — is the build phase ahead.")

    h2(doc, "7.2 From the Solution Architect (hosting, integration, data)")
    qa(doc, "Is this serverless, PaaS, or a VM? And why?",
       "Cloud IaaS — a single right-sized **Azure VM** in a private VNet in Shell's subscription. We chose a VM "
       "for the simplest security review and full control; the app is containerised so we can move to Container "
       "Apps (PaaS) or a Foundry Hosted Agent later without a rewrite. It is **not** on-prem and **not** serverless "
       "today.")
    qa(doc, "How does it talk to Microsoft 365 / our data?",
       "Through **Microsoft Graph**, authenticating app-only with a **certificate**, scoped to a single mailbox via "
       "an Application Access Policy. Database, Key Vault, and Foundry are reached over **Private Endpoints** — "
       "nothing sensitive crosses the public internet.")
    qa(doc, "Where does our data live, and does it leave our tenant?",
       "Everything stays **single-tenant inside Shell's Azure and chosen region**. Foundry runs in Shell's tenant. "
       "Data is not used for model training (confirmed in the DPA). Secrets live in Key Vault, accessed via "
       "Managed Identity — none in code.")
    qa(doc, "How do you handle secrets, identity, and audit?",
       "Secrets in **Key Vault** via **Managed Identity** (no stored passwords). Users authenticate with **Entra "
       "SSO + RBAC**. Every agent run and action is logged with correlation IDs to an immutable **App Insights / "
       "Log Analytics** audit trail.")
    qa(doc, "What about the missing automated tests / tech debt?",
       "Honestly stated: the PoC has no automated suite yet. The production CI pipeline adds a **regression stage** "
       "that gates every deploy and must cover the approval gate, the deterministic path, and the response "
       "contract before any agent change ships.")

    h2(doc, "7.3 From the AI Architect (the agentic layer, safety, model)")
    qa(doc, "Why MAF over just calling the model directly?",
       "MAF gives us platform-built human-in-the-loop approval, content safety, and on-by-default tracing — so we "
       "hand-roll fewer controls, which means a smaller surface for Shell's code-security review. We proved both "
       "the direct-call and the MAF path in the PoC; production builds on MAF, with direct-Responses kept as a "
       "fallback.")
    qa(doc, "How do you stop hallucinations from doing damage?",
       "The model never touches facts, numbers, IDs, or decisions — those are deterministic code. The LLM only "
       "produces draft prose, which is grounded in supplied context and **always** human-approved before use. So a "
       "hallucination is a wording a reviewer rejects, never a wrong action taken.")
    qa(doc, "Which model, and are we locked into Microsoft?",
       "GPT-4o (GA) deployed in Shell's Foundry today. We're **not locked in**: an LLMProvider abstraction keeps "
       "Anthropic/Claude selectable as a configuration choice if procurement prefers it.")
    qa(doc, "How is the agent prevented from sending things on its own?",
       "Two layers. The **app-layer gate**: side-effecting actions (send invite, send mail) are removed from the "
       "agent and only fire from deterministic routes after a human approves. Plus MAF's native **approval_mode** "
       "on any in-run tool. The agent drafts; it never sends.")
    qa(doc, "Do you need GPUs / how heavy is the compute?",
       "No GPUs anywhere. All inference is remote in Foundry. The VM only runs a light FastAPI app, so 4–8 vCPUs "
       "and 16–32 GB RAM is ample.")

    h2(doc, "7.4 Deployment, cost & timeline")
    qa(doc, "What size server and what does it cost to run?",
       "A general-purpose **D4as v6 (4 vCPU / 16 GB)**, up to **D8as v6 (8 vCPU / 32 GB)** in production. Indicative "
       "Azure run-cost is roughly **$450–650/month** for the whole production stack plus **~$1,000/month** of "
       "Foundry model usage — order-of-magnitude, lower with reserved instances. Non-prod adds ~$80–120/month.")
    qa(doc, "What's the biggest thing standing between us and production?",
       "Not the technology — it's **Shell's own compliance gates**: AI Registry + ServiceNow registration, the IRM "
       "risk assessment (IAQ), EU AI Act classification, and Shell.AI + TRB approval, all on IDT-managed Azure. "
       "These have external lead time, so they should start **in parallel, now**.")
    qa(doc, "Can it scale if adoption grows?",
       "Yes — vertically by resizing the VM in minutes, or by moving the containerised app to autoscaling Container "
       "Apps / a Foundry Hosted Agent with no code change. The database is already HA and can scale its tier "
       "independently.")

    # ---- 8. Honest caveats ----
    h1(doc, "8. Honest caveats — so you're never caught out")
    body(doc, "Better you raise these than the client. None are dealbreakers; all have a clean answer.")
    bullet(doc, "**Diagram says App Service icon, doc says VM.** Acknowledge it's a draft artifact; the target is a "
                "VM; both are valid and the code is identical. (§2.1)")
    bullet(doc, "**This is still pre-production.** The PoC proves the pattern; the production build (infra, "
                "identity, DB migration, all agents, tests) is ahead. Don't imply it's deployed.")
    bullet(doc, "**No automated test suite yet** — the CI regression stage is the plan to fix that before any agent "
                "change ships.")
    bullet(doc, "**Foundry Hosted Agents is still preview** — we mention it only as a *future* hosting option; the "
                "shippable baseline is the VM.")
    bullet(doc, "**Compliance gates are blocking and external** — production cannot go live until Shell's AI "
                "Registry / IAQ / TRB / Shell.AI approvals clear. Frame this as a shared workstream to start now, "
                "not a Zensar delay.")
    bullet(doc, "**Costs are indicative** — always say \"order of magnitude, subject to Shell tiering and usage,\" "
                "never quote them as a price.")

    foot = doc.add_paragraph()
    fr = foot.add_run("VendorPulse — Client Meeting Prep Pack · Zensar · Confidential · prepared for the "
                      "prospective-client review")
    fr.font.size = Pt(8)
    fr.font.color.rgb = MUTED
    fr.font.name = "Calibri"
    foot.paragraph_format.space_before = Pt(14)
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
