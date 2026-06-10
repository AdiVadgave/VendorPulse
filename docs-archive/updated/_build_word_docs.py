"""
Build Word (.docx) versions of selected Shell documentation files.

Produces five documents in docs/updated/word/:
  1. 01_Executive_Summary_Shell.docx
  2. 02_Solution_Architecture_Shell.docx
  3. HLD_Backend_and_Frontend_Shell.docx       (combined 03 + 04)
  4. 08_Dependencies_and_Access_Requirements.docx
  5. 11_Productionization_Roadmap_Shell.docx

Pleasant readable theme:
  - H1 deep navy, H2 medium blue, H3 teal, H4 slate
  - Body Calibri 11pt slate grey
  - Tables with soft blue header + alternating tints
  - Code with light grey background + monospace
  - Callouts (blockquotes) with gold left border
  - Page numbers in footer

Run from any directory; resolves paths relative to this script.
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches

# ---------------------------------------------------------------------------
# Theme
# ---------------------------------------------------------------------------

C_NAVY        = RGBColor(0x1B, 0x3A, 0x6F)   # H1
C_BLUE        = RGBColor(0x0F, 0x5B, 0xA8)   # H2
C_TEAL        = RGBColor(0x0B, 0x80, 0x8C)   # H3
C_SLATE       = RGBColor(0x33, 0x40, 0x55)   # H4
C_BODY        = RGBColor(0x1E, 0x29, 0x3B)   # body
C_MUTED       = RGBColor(0x64, 0x74, 0x8B)   # captions
C_LINK        = RGBColor(0x0F, 0x5B, 0xA8)
C_ACCENT_GOLD = RGBColor(0xC9, 0x9A, 0x06)   # callout border
C_CODE_TXT    = RGBColor(0x33, 0x40, 0x55)

# Fills (hex without #)
F_TABLE_HEAD  = "DCE7F4"   # soft blue
F_TABLE_ALT   = "F4F7FB"   # very light blue
F_CODE        = "F1F3F6"   # light grey
F_CALLOUT     = "FFF8E7"   # very pale gold
F_RULE        = "C7D0DD"   # divider

FONT_BODY = "Calibri"
FONT_MONO = "Consolas"

H_FONT = "Calibri"

DOCS_DIR  = Path(__file__).resolve().parent
OUT_DIR   = DOCS_DIR / "word"
OUT_DIR.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def shd(color_hex: str):
    """Return a shading element for a cell or paragraph."""
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), color_hex)
    return el


def set_cell_shading(cell, color_hex: str):
    cell._tc.get_or_add_tcPr().append(shd(color_hex))


def set_paragraph_shading(paragraph, color_hex: str):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(shd(color_hex))


def set_cell_borders(cell, color_hex="C7D0DD", size=4):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:color"), color_hex)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def left_border(paragraph, color: str, size=24):
    """Add a thick left border to a paragraph for callouts."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:left")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size))
    b.set(qn("w:space"), "8")
    b.set(qn("w:color"), color)
    pBdr.append(b)
    pPr.append(pBdr)


def horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), F_RULE)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


# ---------------------------------------------------------------------------
# Document scaffolding
# ---------------------------------------------------------------------------

def new_doc(title: str, subtitle: str = "") -> Document:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.2)

    # Base style
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(11)
    style.font.color.rgb = C_BODY

    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    # Title block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run(title)
    title_run.font.name = H_FONT
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = C_NAVY

    if subtitle:
        sub_p = doc.add_paragraph()
        sub_run = sub_p.add_run(subtitle)
        sub_run.font.name = FONT_BODY
        sub_run.font.size = Pt(12)
        sub_run.font.italic = True
        sub_run.font.color.rgb = C_MUTED

    horizontal_rule(doc)
    add_footer_page_number(doc)
    return doc


def add_footer_page_number(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.font.name = FONT_BODY
    run.font.size = Pt(9)
    run.font.color.rgb = C_MUTED

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


# ---------------------------------------------------------------------------
# Inline formatting (bold/italic/code/links)
# ---------------------------------------------------------------------------

INLINE_RE = re.compile(
    r"(\*\*[^*]+?\*\*)"           # **bold**
    r"|(\*[^*]+?\*)"              # *italic*
    r"|(`[^`]+?`)"                # `code`
    r"|(\[[^\]]+?\]\([^)]+?\))"   # [link](url)
)


def render_inline(paragraph, text: str, base_color=C_BODY, base_size=11):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            r.font.name = FONT_BODY
            r.font.size = Pt(base_size)
            r.font.color.rgb = base_color

        token = m.group(0)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.font.bold = True
            r.font.name = FONT_BODY
            r.font.size = Pt(base_size)
            r.font.color.rgb = base_color
        elif token.startswith("*"):
            r = paragraph.add_run(token[1:-1])
            r.font.italic = True
            r.font.name = FONT_BODY
            r.font.size = Pt(base_size)
            r.font.color.rgb = base_color
        elif token.startswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = FONT_MONO
            r.font.size = Pt(base_size - 1)
            r.font.color.rgb = C_CODE_TXT
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), F_CODE)
            r._r.get_or_add_rPr().append(shading)
        elif token.startswith("["):
            # link → render text only, styled as link
            mm = re.match(r"\[([^\]]+?)\]\(([^)]+?)\)", token)
            link_text = mm.group(1) if mm else token
            r = paragraph.add_run(link_text)
            r.font.color.rgb = C_LINK
            r.font.underline = True
            r.font.name = FONT_BODY
            r.font.size = Pt(base_size)
        pos = m.end()

    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.font.name = FONT_BODY
        r.font.size = Pt(base_size)
        r.font.color.rgb = base_color


# ---------------------------------------------------------------------------
# Block elements
# ---------------------------------------------------------------------------

def add_heading_styled(doc, text, level):
    sizes = {1: 22, 2: 16, 3: 13, 4: 12}
    colors = {1: C_NAVY, 2: C_BLUE, 3: C_TEAL, 4: C_SLATE}
    spacing_before = {1: 18, 2: 14, 3: 10, 4: 8}
    spacing_after = {1: 8, 2: 6, 3: 4, 4: 3}

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spacing_before.get(level, 6))
    p.paragraph_format.space_after = Pt(spacing_after.get(level, 4))
    p.paragraph_format.keep_with_next = True

    run = p.add_run(text)
    run.font.name = H_FONT
    run.font.size = Pt(sizes.get(level, 12))
    run.font.bold = True
    run.font.color.rgb = colors.get(level, C_BODY)


def add_callout(doc, lines: list[str]):
    """Blockquote → callout with light-gold background + gold left border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.4)
    set_paragraph_shading(p, F_CALLOUT)
    left_border(p, "C99A06", size=24)

    text = " ".join(l.lstrip("> ").strip() for l in lines)
    render_inline(p, text)


def add_code_block(doc, lines: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.2)
    set_paragraph_shading(p, F_CODE)
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line.rstrip())
        r.font.name = FONT_MONO
        r.font.size = Pt(9.5)
        r.font.color.rgb = C_CODE_TXT


def add_list_item(doc, text: str, ordered: bool, level: int = 0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6 + 0.6 * level)
    p.paragraph_format.space_after = Pt(3)
    bullet = "• " if not ordered else ""
    if not ordered:
        r = p.add_run(bullet)
        r.font.color.rgb = C_BLUE
        r.font.bold = True
    render_inline(p, text)


def add_table(doc, rows: list[list[str]]):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    # pad short rows
    for r in rows:
        while len(r) < n_cols:
            r.append("")

    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = True

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = tbl.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_borders(cell, color_hex="C7D0DD", size=4)

            if i == 0:
                set_cell_shading(cell, F_TABLE_HEAD)
            elif i % 2 == 0:
                set_cell_shading(cell, F_TABLE_ALT)

            # write content
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)

            if i == 0:
                # header cell
                r = para.add_run(cell_text.strip())
                r.font.bold = True
                r.font.color.rgb = C_NAVY
                r.font.name = FONT_BODY
                r.font.size = Pt(10.5)
            else:
                render_inline(para, cell_text.strip(), base_size=10)

    doc.add_paragraph()  # spacer after table


# ---------------------------------------------------------------------------
# Markdown parsing
# ---------------------------------------------------------------------------

HEADING_RE   = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")
LIST_UL_RE   = re.compile(r"^(\s*)[-*+]\s+(.*)$")
LIST_OL_RE   = re.compile(r"^(\s*)\d+\.\s+(.*)$")
HR_RE        = re.compile(r"^\s*-{3,}\s*$")


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def render_markdown(doc, md: str):
    """Parse a subset of markdown and write into the Word doc."""
    lines = md.splitlines()
    i = 0
    in_frontmatter = False
    while i < len(lines):
        line = lines[i]

        # Skip YAML front-matter if any
        if i == 0 and line.strip() == "---":
            in_frontmatter = True
            i += 1
            continue
        if in_frontmatter:
            if line.strip() == "---":
                in_frontmatter = False
            i += 1
            continue

        # Code fence
        if line.lstrip().startswith("```"):
            j = i + 1
            block = []
            while j < len(lines) and not lines[j].lstrip().startswith("```"):
                block.append(lines[j])
                j += 1
            add_code_block(doc, block)
            i = j + 1
            continue

        # Horizontal rule
        if HR_RE.match(line):
            horizontal_rule(doc)
            i += 1
            continue

        # Heading
        m = HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # Strip any inline-link wrapping that's pure
            text = re.sub(r"^\[([^\]]+?)\]\([^)]+?\)$", r"\1", text)
            add_heading_styled(doc, text, level)
            i += 1
            continue

        # Blockquote / callout
        if line.lstrip().startswith(">"):
            j = i
            block = []
            while j < len(lines) and lines[j].lstrip().startswith(">"):
                block.append(lines[j])
                j += 1
            add_callout(doc, block)
            i = j
            continue

        # Table (must have at least 2 lines and second is the separator)
        if "|" in line and i + 1 < len(lines) and TABLE_SEP_RE.match(lines[i + 1]):
            header = split_table_row(line)
            j = i + 2
            rows = [header]
            while j < len(lines) and "|" in lines[j] and lines[j].strip():
                rows.append(split_table_row(lines[j]))
                j += 1
            add_table(doc, rows)
            i = j
            continue

        # Unordered list
        m = LIST_UL_RE.match(line)
        if m:
            indent = len(m.group(1))
            level = indent // 2
            add_list_item(doc, m.group(2).strip(), ordered=False, level=level)
            i += 1
            continue

        # Ordered list
        m = LIST_OL_RE.match(line)
        if m:
            indent = len(m.group(1))
            level = indent // 2
            add_list_item(doc, m.group(2).strip(), ordered=True, level=level)
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Paragraph (may span multiple consecutive non-empty lines)
        para_lines = [line]
        j = i + 1
        while (
            j < len(lines)
            and lines[j].strip()
            and not HEADING_RE.match(lines[j])
            and not lines[j].lstrip().startswith(("```", ">", "-", "*", "+", "|"))
            and not LIST_OL_RE.match(lines[j])
            and not HR_RE.match(lines[j])
        ):
            para_lines.append(lines[j])
            j += 1

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        render_inline(p, " ".join(l.strip() for l in para_lines))
        i = j


# ---------------------------------------------------------------------------
# Per-document recipes
# ---------------------------------------------------------------------------

@dataclass
class DocRecipe:
    title: str
    subtitle: str
    sources: list[Path]
    output: Path
    section_break_between: bool = True


def build_doc(recipe: DocRecipe):
    doc = new_doc(recipe.title, recipe.subtitle)

    for idx, src in enumerate(recipe.sources):
        if idx > 0 and recipe.section_break_between:
            add_page_break(doc)
        md = src.read_text(encoding="utf-8")
        render_markdown(doc, md)

    doc.save(recipe.output)
    print(f"  wrote {recipe.output.relative_to(DOCS_DIR)}")


def main():
    recipes = [
        DocRecipe(
            title="VendorPulse — Executive Summary",
            subtitle="Shell engagement · Version 2.0 · 2026-06-03",
            sources=[DOCS_DIR / "01_Executive_Summary_Shell.md"],
            output=OUT_DIR / "01_Executive_Summary_Shell.docx",
        ),
        DocRecipe(
            title="VendorPulse — Solution Architecture",
            subtitle="Shell engagement · Version 2.0 · 2026-06-03",
            sources=[DOCS_DIR / "02_Solution_Architecture_Shell.md"],
            output=OUT_DIR / "02_Solution_Architecture_Shell.docx",
        ),
        DocRecipe(
            title="VendorPulse — HLD (Backend + Frontend)",
            subtitle="Combined High-Level Design · Shell engagement · Version 2.0",
            sources=[
                DOCS_DIR / "03_HLD_Backend_Shell.md",
                DOCS_DIR / "04_HLD_Frontend_Shell.md",
            ],
            output=OUT_DIR / "HLD_Backend_and_Frontend_Shell.docx",
        ),
        DocRecipe(
            title="VendorPulse — Dependencies & Access Requirements",
            subtitle="Shell engagement · Version 2.0 · 2026-06-03",
            sources=[DOCS_DIR / "08_Dependencies_and_Access_Requirements.md"],
            output=OUT_DIR / "08_Dependencies_and_Access_Requirements.docx",
        ),
        DocRecipe(
            title="VendorPulse — Productionization Roadmap (3-Week Plan)",
            subtitle="Shell engagement · Version 2.2 · 2026-06-03",
            sources=[DOCS_DIR / "11_Productionization_Roadmap_Shell.md"],
            output=OUT_DIR / "11_Productionization_Roadmap_Shell.docx",
        ),
    ]

    print(f"Building {len(recipes)} Word documents -> {OUT_DIR}")
    for r in recipes:
        build_doc(r)
    print("Done.")


if __name__ == "__main__":
    main()
