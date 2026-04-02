"""
VendorPulse MVP Development Plan - DOCX Generator
Generates a professionally formatted Word document.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── COLOUR PALETTE ──────────────────────────────────────────────────────────
# Stored as (r, g, b) tuples; use rgb() helper for RGBColor and hex() for XML
_NAVY        = (0x00, 0x2D, 0x5C)
_BLUE        = (0x00, 0x63, 0xB1)
_LIGHT_BLUE  = (0xDA, 0xE8, 0xF5)
_TEAL        = (0x00, 0x7A, 0x87)
_GOLD        = (0xC9, 0x9A, 0x06)
_WHITE       = (0xFF, 0xFF, 0xFF)
_DARK_GREY   = (0x1A, 0x1A, 0x2E)
_MID_GREY    = (0x55, 0x55, 0x55)
_LIGHT_GREY  = (0xF4, 0xF6, 0xF9)
_BORDER_GREY = (0xBF, 0xBF, 0xBF)

def rgb(t):
    """Convert (r,g,b) tuple to RGBColor."""
    return RGBColor(t[0], t[1], t[2])

def hexc(t):
    """Convert (r,g,b) tuple to uppercase hex string for XML."""
    return '{:02X}{:02X}{:02X}'.format(t[0], t[1], t[2])

NAVY        = _NAVY
BLUE        = _BLUE
LIGHT_BLUE  = _LIGHT_BLUE
TEAL        = _TEAL
GOLD        = _GOLD
WHITE       = _WHITE
DARK_GREY   = _DARK_GREY
MID_GREY    = _MID_GREY
LIGHT_GREY  = _LIGHT_GREY
BORDER_GREY = _BORDER_GREY


# ─── HELPERS ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, color_tuple):
    """Fill a table cell with a solid background colour. Accepts (r,g,b) tuple."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexc(color_tuple))
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), val.get('val', 'single'))
            b.set(qn('w:sz'), str(val.get('sz', 4)))
            b.set(qn('w:color'), val.get('color', 'BFBFBF'))
            tcBorders.append(b)
    tcPr.append(tcBorders)


def set_table_border(table):
    """Add subtle border around and between all cells."""
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={'val': 'single', 'sz': 4, 'color': 'BFBFBF'},
                bottom={'val': 'single', 'sz': 4, 'color': 'BFBFBF'},
                left={'val': 'single', 'sz': 4, 'color': 'BFBFBF'},
                right={'val': 'single', 'sz': 4, 'color': 'BFBFBF'},
            )


def add_run(para, text, bold=False, italic=False, color=None, size=None, font_name=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if font_name:
        run.font.name = font_name
    return run


def add_heading1(doc, text):
    """Top-level section heading — navy background bar."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after = Pt(6)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), '002D5C')
    para._p.get_or_add_pPr().append(shading)
    para.paragraph_format.left_indent = Inches(0.1)
    run = para.add_run('  ' + text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = rgb(WHITE)
    run.font.name = 'Calibri'
    return para


def add_heading2(doc, text):
    """Sub-section heading — blue left border style."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)
    # Blue underline via bottom border on paragraph
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '0063B1')
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = rgb(NAVY)
    run.font.name = 'Calibri'
    return para


def add_heading3(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEAL
    run.font.name = 'Calibri'
    return para


def add_body(doc, text, indent=0):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(4)
    if indent:
        para.paragraph_format.left_indent = Inches(indent)
    run = para.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_GREY
    run.font.name = 'Calibri'
    return para


def add_bullet(doc, text, level=0, bold_prefix=None):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    try:
        para = doc.add_paragraph(style=style)
    except Exception:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = para.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = DARK_GREY
        r1.font.name = 'Calibri'
        r2 = para.add_run(text)
        r2.font.size = Pt(10)
        r2.font.color.rgb = DARK_GREY
        r2.font.name = 'Calibri'
    else:
        run = para.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GREY
        run.font.name = 'Calibri'
    return para


def add_code_block(doc, text):
    """Monospaced grey background code/pre block."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.left_indent = Inches(0.25)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F6F9')
    pPr.append(shd)
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1F, 0x50, 0x7D)
    return para


def add_info_box(doc, label, text):
    """Teal-left-bordered info callout."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.left_indent = Inches(0.3)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EAF7F8')
    pPr.append(shd)
    r1 = para.add_run(label + '  ')
    r1.bold = True
    r1.font.color.rgb = TEAL
    r1.font.size = Pt(10)
    r1.font.name = 'Calibri'
    r2 = para.add_run(text)
    r2.font.color.rgb = DARK_GREY
    r2.font.size = Pt(10)
    r2.font.name = 'Calibri'


def add_standard_table(doc, headers, rows, col_widths=None):
    """
    headers: list of strings
    rows: list of lists (each inner list = one row, same length as headers)
    col_widths: list of Inches values (optional)
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, NAVY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9.5)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = LIGHT_BLUE if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            # Support inline bold via ** markers
            parts = val.split('**')
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                run.bold = (idx % 2 == 1)
                run.font.size = Pt(9.5)
                run.font.color.rgb = DARK_GREY
                run.font.name = 'Calibri'

    # Column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = w

    set_table_border(table)
    doc.add_paragraph()  # spacing after table
    return table


def add_workflow_state_table(doc, states):
    """Special coloured workflow states table."""
    table = doc.add_table(rows=len(states), cols=2)
    table.style = 'Table Grid'
    colors = [
        RGBColor(0x00, 0x2D, 0x5C),
        RGBColor(0x00, 0x50, 0x9A),
        RGBColor(0x00, 0x70, 0xC0),
        RGBColor(0x00, 0x7A, 0x87),
        RGBColor(0x2E, 0x86, 0x5A),
        RGBColor(0x5C, 0x94, 0x30),
        RGBColor(0xC9, 0x9A, 0x06),
        RGBColor(0xC0, 0x50, 0x00),
        RGBColor(0x80, 0x27, 0x00),
        RGBColor(0x55, 0x55, 0x55),
        RGBColor(0x33, 0x33, 0x55),
        RGBColor(0x1A, 0x1A, 0x3A),
        RGBColor(0x0D, 0x0D, 0x20),
        RGBColor(0x4B, 0x0F, 0x6F),
        RGBColor(0x35, 0x0A, 0x4F),
    ]
    for i, (state, trigger) in enumerate(states):
        c = colors[i % len(colors)]
        cell0, cell1 = table.rows[i].cells[0], table.rows[i].cells[1]
        set_cell_bg(cell0, c)
        p0 = cell0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(3)
        p0.paragraph_format.space_after = Pt(3)
        r = p0.add_run(state)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(9)
        r.font.name = 'Courier New'

        set_cell_bg(cell1, RGBColor(0xF8, 0xF9, 0xFA))
        p1 = cell1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(3)
        p1.paragraph_format.space_after = Pt(3)
        r2 = p1.add_run(trigger)
        r2.font.size = Pt(9)
        r2.font.color.rgb = DARK_GREY
        r2.font.name = 'Calibri'

    set_table_border(table)
    doc.add_paragraph()
    return table


def add_phase_box(doc, day_label, title, bullets):
    """Coloured phase/day box."""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    phase_colors = {
        'Day 1': RGBColor(0x00, 0x2D, 0x5C),
        'Day 2': RGBColor(0x00, 0x63, 0xB1),
        'Day 3': RGBColor(0x00, 0x7A, 0x87),
        'Day 4': RGBColor(0x2E, 0x86, 0x5A),
        'Day 5': RGBColor(0xC9, 0x9A, 0x06),
    }
    bg = phase_colors.get(day_label, NAVY)

    # Left column — day label
    cell0 = table.rows[0].cells[0]
    set_cell_bg(cell0, bg)
    cell0.width = Inches(0.9)
    p0 = cell0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(6)
    p0.paragraph_format.space_after = Pt(2)
    r = p0.add_run(day_label)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(12)
    r.font.name = 'Calibri'
    p0.add_run('\n')
    r2 = p0.add_run(title)
    r2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
    r2.font.size = Pt(8.5)
    r2.font.name = 'Calibri'

    # Right column — bullets
    cell1 = table.rows[0].cells[1]
    set_cell_bg(cell1, RGBColor(0xF4, 0xF6, 0xF9))
    for b in bullets:
        p = cell1.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.15)
        run = p.add_run('• ' + b)
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK_GREY
        run.font.name = 'Calibri'

    set_table_border(table)
    doc.add_paragraph()


def add_page_break(doc):
    doc.add_page_break()


def set_doc_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


# ─── COVER PAGE ───────────────────────────────────────────────────────────────

def add_cover_page(doc):
    # Big navy header block
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(40)
    para.paragraph_format.space_after = Pt(0)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '002D5C')
    pPr.append(shd)
    para.paragraph_format.left_indent = Inches(0.3)
    r1 = para.add_run('  VendorPulse')
    r1.bold = True
    r1.font.size = Pt(32)
    r1.font.color.rgb = WHITE
    r1.font.name = 'Calibri'

    para2 = doc.add_paragraph()
    para2.paragraph_format.space_before = Pt(0)
    para2.paragraph_format.space_after = Pt(0)
    pPr2 = para2._p.get_or_add_pPr()
    shd2 = OxmlElement('w:shd')
    shd2.set(qn('w:val'), 'clear')
    shd2.set(qn('w:color'), 'auto')
    shd2.set(qn('w:fill'), '0063B1')
    pPr2.append(shd2)
    para2.paragraph_format.left_indent = Inches(0.3)
    r2 = para2.add_run('  MVP Development Plan  ·  v2.0')
    r2.bold = True
    r2.font.size = Pt(16)
    r2.font.color.rgb = WHITE
    r2.font.name = 'Calibri'

    doc.add_paragraph()

    meta = [
        ('Product', 'Agentic AI for Governance Cycle Automation'),
        ('Client', 'Shell EGB/QBR Governance Demo'),
        ('Prepared by', 'Zensar Technologies'),
        ('Sprint Duration', '5 Working Days'),
        ('Frontend', 'React 18 + Vite + Tailwind CSS + shadcn/ui'),
        ('Backend', 'FastAPI (Python) + SQLite'),
        ('AI Engine', 'Anthropic Claude API — Tool-Calling Pattern'),
        ('Status', 'Draft — For Internal Team Alignment'),
        ('Date', '2026-04-01'),
    ]
    add_standard_table(doc, ['Field', 'Value'], meta, [Inches(2.0), Inches(4.0)])

    doc.add_paragraph()
    add_info_box(doc, 'What VendorPulse Is:',
        'A workflow orchestrator · A document intelligence layer · '
        'A governance memory system · A human-approved automation assistant')


# ─── DOCUMENT BUILDER ────────────────────────────────────────────────────────

def build_document():
    doc = Document()
    set_doc_margins(doc)

    # Default paragraph font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # ── COVER ──────────────────────────────────────────────────────────────
    add_cover_page(doc)
    add_page_break(doc)

    # ── SECTION 1: SYSTEM OVERVIEW ─────────────────────────────────────────
    add_heading1(doc, '1.  System Overview & Goals')
    add_body(doc,
        'VendorPulse automates Shell\'s EGB (Executive Governance Board) and QBR (Quarterly '
        'Business Review) governance cycles end-to-end — from scheduling through post-meeting '
        'analytics. It is not a chatbot.')
    add_bullet(doc, 'Workflow orchestrator — moves a governance cycle through defined stages')
    add_bullet(doc, 'Document intelligence layer — generates briefs, minutes and responses via Claude')
    add_bullet(doc, 'Governance memory system — stores historical data across cycles')
    add_bullet(doc, 'Human-approved automation assistant — AI suggests, humans approve')

    add_heading2(doc, 'Problem Statement Summary')
    add_standard_table(doc,
        ['#', 'Pain Point', 'Observed Impact'],
        [
            ['1', 'High administrative overhead scheduling multi-stakeholder meetings', '3–5 hrs per cycle per coordinator'],
            ['2', 'Manual consolidation and validation of scorecard inputs', 'Error-prone, delayed reporting'],
            ['3', 'Frequent manual reminders to Shell stakeholders and vendors', 'Chase email fatigue'],
            ['4', 'Vendor pushback requires multiple back-and-forth cycles', 'Governance delays'],
            ['5', 'Limited cross-cycle memory on trends and prior agreements', 'Missed escalation opportunities'],
            ['6', 'Preparation work outweighs value-driven governance work', 'Strategic capacity lost'],
        ],
        [Inches(0.3), Inches(3.2), Inches(2.5)]
    )

    add_heading2(doc, 'Prototype Scope')
    add_standard_table(doc,
        ['Module', 'Name', 'In Scope v1.0'],
        [
            ['A', 'Meeting Scheduling & Coordination Agent', 'Yes'],
            ['B', 'Scorecard Input Collection & Validation Agent', 'Yes'],
            ['C', 'Internal Alignment Call Support Agent', 'Yes'],
            ['D', 'Vendor Prep Call Support Agent', 'Yes'],
            ['E', 'EGB/QBR Live Meeting Support Agent', 'Yes'],
            ['F', 'Cross-Cycle Memory & Trend Analysis Dashboard', 'Yes'],
            ['—', 'Contract Management Automation (Section B)', 'Phase 2 only'],
            ['—', 'Real calendar / email integration', 'Mock APIs — Phase 2'],
            ['—', 'PPT / slide auto-population', 'Out of scope'],
        ],
        [Inches(0.7), Inches(3.5), Inches(1.8)]
    )

    add_heading2(doc, 'User Personas')
    add_standard_table(doc,
        ['Persona', 'Role', 'Primary Modules'],
        [
            ['Alex — VMO Coordinator', 'Manages scheduling, reminders, scorecard chasing', 'A, B'],
            ['Priya — Internal Lead', 'Prepares internal alignment, reviews vendor scores', 'C, F'],
            ['Marcus — Vendor Manager', 'Manages vendor pre-call, drafts pushback responses', 'D'],
            ['Sandra — EGB Chair', 'Runs live governance meetings, reviews minutes', 'E, F'],
        ],
        [Inches(2.0), Inches(3.0), Inches(1.0)]
    )
    add_page_break(doc)

    # ── SECTION 2: TECH STACK ──────────────────────────────────────────────
    add_heading1(doc, '2.  Technology Stack')
    add_standard_table(doc,
        ['Layer', 'Technology', 'Purpose'],
        [
            ['Frontend Framework', 'React 18 + Vite', 'SPA with fast HMR dev experience'],
            ['UI Styling', 'Tailwind CSS + shadcn/ui', 'Utility-first + accessible component primitives'],
            ['Charts', 'Recharts', 'Line, Radar, Bar charts for Module F'],
            ['Forms', 'React Hook Form + Zod', 'Typed validation without re-renders'],
            ['Client State', 'Zustand', 'Lightweight global store (cycle, approvals)'],
            ['Server State', 'TanStack Query v5', 'Caching, background refetch, loading states'],
            ['Backend', 'FastAPI (Python 3.11+)', 'Async REST API with automatic OpenAPI docs'],
            ['ORM', 'SQLAlchemy 2.0 (async)', 'Typed DB access, easy migration to Postgres'],
            ['Database', 'SQLite (file-based)', 'Zero-config, demo-reset friendly'],
            ['Migrations', 'Alembic', 'Version-controlled schema changes'],
            ['Validation', 'Pydantic v2', 'Request/response schema enforcement'],
            ['AI Engine', 'Anthropic Claude API', 'Tool-calling pattern per agent module'],
            ['Mock Services', 'In-process Python classes', 'Calendar, Email, Forms, Notifications'],
        ],
        [Inches(2.0), Inches(2.2), Inches(2.0)]
    )
    add_page_break(doc)

    # ── SECTION 3: PROJECT STRUCTURE ───────────────────────────────────────
    add_heading1(doc, '3.  Project Folder Structure')

    add_heading2(doc, 'Backend (FastAPI)')
    add_code_block(doc, """backend/
  app/
    main.py
    api/
      routes/
        cycles.py          ← cycle CRUD & workflow state
        scheduling.py      ← Module A endpoints
        scorecard.py       ← Module B endpoints
        alignment.py       ← Module C endpoints
        vendor_prep.py     ← Module D endpoints
        meeting.py         ← Module E endpoints
        analytics.py       ← Module F endpoints
    core/
      config.py
      database.py
      workflow_engine.py   ← state machine
    agents/
      base_agent.py        ← shared tool-calling pattern
      scheduling_agent.py
      scorecard_agent.py
      alignment_agent.py
      vendor_prep_agent.py
      meeting_agent.py
      memory_agent.py
    services/
      mock/
        mock_calendar.py
        mock_email.py
        mock_forms.py
        mock_notifications.py
      llm_service.py       ← Claude API wrapper
      validation_service.py
      analytics_service.py
    models/                ← SQLAlchemy ORM models
    schemas/               ← Pydantic schemas (request/response)
    repositories/          ← DB access layer
    utils/
      prompts.py           ← all LLM prompt templates
      slot_ranking.py      ← deterministic slot ranker
      score_diff.py        ← cycle comparison engine
      text_parsing.py      ← action item extractor
  seed/
    seed_data.py
  alembic/
  tests/""")

    add_heading2(doc, 'Frontend (React + Vite)')
    add_code_block(doc, """frontend/
  src/
    components/
      ui/                  ← shadcn/ui primitives (Button, Card, Badge, Dialog…)
      modules/
        scheduling/        ← AttendeeRefreshPanel, SlotRankingPanel, ConfirmationTracker
        scorecard/         ← DispatchPanel, SubmissionTracker, CompiledScorecardTable
        alignment/         ← ChangeHighlights, AlignmentFlags, FaceOffModelEditor
        vendor-prep/       ← VendorBriefPanel, PushbackInput, ResponseCards
        meeting/           ← LiveCapturePanel, MeetingMinutesViewer
        analytics/         ← TrendLineChart, RadarChart, RecurringIssueAlerts
      shared/
        ApprovalPanel.tsx  ← reused before every "send" action
        ActionLog.tsx      ← reused across Modules C, D, E
        AgentStatusBadge.tsx
        WorkflowProgressBar.tsx
    pages/
      Dashboard.tsx        ← active cycles overview
      CycleDetail.tsx      ← main tabbed workspace per cycle
      Analytics.tsx        ← Module F charts & briefings
    api/                   ← typed axios/fetch wrappers per module
    store/                 ← Zustand stores
    hooks/                 ← custom React hooks per module
    types/                 ← TypeScript types matching backend schemas""")

    add_page_break(doc)

    # ── SECTION 4: DATA MODEL ──────────────────────────────────────────────
    add_heading1(doc, '4.  Data Model')
    add_body(doc, 'All 13 tables. Every entity is tied to a Governance Cycle as the central organising concept.')

    tables_data = [
        ('vendors', 'vendor_id · name · category · status · created_at',
         'Seed: NovaTech Services, CoreSystems Ltd, Meridian IT'),
        ('governance_cycles', 'cycle_id · vendor_id · cycle_name · quarter · year · **workflow_state** · created_at · updated_at',
         'workflow_state drives the entire product flow'),
        ('stakeholders', 'stakeholder_id · name · email · **role** · organisation · is_active · created_at',
         'Roles: VMO_COORDINATOR, INTERNAL_LEAD, VENDOR_MANAGER, EGB_CHAIR, TECHNICAL_LEAD, COMMERCIAL_LEAD'),
        ('cycle_attendees', 'id · cycle_id · stakeholder_id · is_confirmed · **is_key** · replacement_name · replacement_email · invite_status',
         'is_key=true for organiser & exec sponsor (hard constraint in slot ranking)'),
        ('scorecards', 'scorecard_id · cycle_id · stakeholder_id · vendor_id · **category** · score(1-5) · comment · is_valid · validation_flags(JSON) · submitted_at',
         'Categories: DELIVERY_QUALITY, SLA_COMPLIANCE, INNOVATION, COMMUNICATION, VALUE_FOR_MONEY'),
        ('meetings', 'meeting_id · cycle_id · **meeting_type** · scheduled_time · location_or_dial_in · invite_sent_at · minutes_generated_at',
         'Types: INTERNAL_ALIGNMENT, VENDOR_PREP, EGB_QBR'),
        ('meeting_notes', 'note_id · meeting_id · **note_type** · content · raised_by_role · timestamp · is_actioned',
         'Types: QUESTION, OBJECTION, DECISION, APPRECIATION, ACTION'),
        ('action_items', 'action_id · cycle_id · **source_module** · description · owner · due_date · status · created_at',
         'source_module: ALIGNMENT | VENDOR_PREP | MEETING — merged in unified action log'),
        ('issues', 'issue_id · vendor_id · description · first_seen_cycle_id · **occurrences** · status · last_owner · last_updated',
         'occurrences >= 2 triggers recurring issue alert in Module F'),
        ('face_off_model', 'id · cycle_id · position_number · shell_name · shell_role · vendor_name · vendor_role',
         'Updated via inline form in Modules C and D'),
        ('notifications', 'notification_id · cycle_id · stakeholder_id · **type** · content · sent_at · status',
         'Types: SCORECARD_REQUEST, REMINDER_1, REMINDER_2, ESCALATION, INVITE'),
        ('slot_proposals', 'slot_id · cycle_id · proposed_time · timezone · organiser_available · exec_sponsor_available · attendee_availability(JSON) · rank_score · is_approved',
         'rank_score calculated deterministically — no LLM'),
        ('agent_runs', 'run_id · agent_name · cycle_id · input_payload(JSON) · output_payload(JSON) · status · error_message · triggered_by · created_at',
         'Every agent action logged here — critical for traceability and debugging'),
    ]

    add_standard_table(doc,
        ['Table', 'Key Columns', 'Notes'],
        tables_data,
        [Inches(1.4), Inches(2.8), Inches(2.0)]
    )
    add_page_break(doc)

    # ── SECTION 5: AGENT RESPONSE CONTRACT ────────────────────────────────
    add_heading1(doc, '5.  Standard Agent Response Contract')
    add_body(doc,
        'Every agent returns this exact shape. The frontend never parses raw AI text. '
        'This contract is enforced by a Pydantic schema on the backend.')
    add_code_block(doc, """{
  "status":           "success | failed | partial | pending_approval",
  "agent":            "scheduling_agent",
  "summary":          "Attendee refresh form generated for 9 stakeholders.",
  "data":             { ... },
  "warnings":         ["3 stakeholders have not responded"],
  "next_actions":     ["APPROVE_INVITE", "SEND_REMINDER"],
  "requires_approval": true,
  "run_id":           "uuid"
}""")
    add_info_box(doc, 'Why this matters:',
        'The UI never guesses shape. Approval flows are driven by requires_approval. '
        'Warnings surface without blocking. next_actions auto-highlights the next button. '
        'run_id links back to the agent_runs table for full traceability.')

    add_page_break(doc)

    # ── SECTION 6: WORKFLOW ENGINE ─────────────────────────────────────────
    add_heading1(doc, '6.  Workflow Engine — State Machine')
    add_body(doc,
        'The workflow_state on each cycle only moves forward and only when prerequisites are met. '
        'This is pure deterministic logic — no LLM involvement.')

    states = [
        ('CYCLE_CREATED',              'Trigger: user clicks "Start New Cycle"'),
        ('ATTENDEE_REFRESH_SENT',      'Trigger: organiser approves refresh form dispatch'),
        ('AVAILABILITY_COLLECTED',     'Trigger: all key attendees have responded'),
        ('MEETING_SCHEDULED',          'Trigger: organiser approves a slot'),
        ('SCORECARD_REQUEST_SENT',     'Trigger: organiser approves scorecard dispatch'),
        ('SCORECARD_COLLECTION',       'Trigger: at least 1 submission received'),
        ('SCORECARD_COMPILED',         'Trigger: deadline passed OR manual compile'),
        ('INTERNAL_ALIGNMENT',         'Trigger: scorecard compiled'),
        ('VENDOR_PREP',                'Trigger: alignment notes saved'),
        ('MEETING_IN_PROGRESS',        'Trigger: facilitator clicks "Start Meeting"'),
        ('POST_MEETING_COMPLETE',      'Trigger: minutes approved'),
        ('ARCHIVED',                   'Trigger: manual — all open actions closed'),
    ]
    add_workflow_state_table(doc, states)

    add_heading2(doc, 'Enforcement Rules (examples)')
    add_bullet(doc, 'Cannot send vendor brief before scorecard is compiled')
    add_bullet(doc, 'Cannot generate meeting minutes before meeting notes exist')
    add_bullet(doc, 'Cannot move to ARCHIVED without at least one approved action item log')
    add_bullet(doc, 'Cannot compile scorecard before at least 2 valid submissions are received')
    add_page_break(doc)

    # ── SECTION 7: MOCK SERVICES ───────────────────────────────────────────
    add_heading1(doc, '7.  Mock Services Layer')
    add_body(doc,
        'All four services implement a clean interface so they can be swapped for real '
        'Outlook / Teams / SharePoint integrations later without rewriting agent code.')

    add_standard_table(doc,
        ['Service', 'Key Method', 'What It Does in Demo'],
        [
            ['**MockCalendarService**', 'get_availability(stakeholder_ids, date_range)', 'Returns fixture availability from seeded schedule data — no external calls'],
            ['**MockEmailService**', 'send(to, subject, body)', 'Stores "sent" emails to mock_outbox; returns HTML preview for approval panel'],
            ['**MockFormService**', 'create_form(type, fields, recipients)', 'Opens form as in-app modal; simulates async responses via seed data + "Simulate Responses" button'],
            ['**MockNotificationService**', 'send_reminder(stakeholder_id, level)', 'Writes to notifications table; rendered in Notifications panel with escalating tone labels'],
        ],
        [Inches(1.8), Inches(2.6), Inches(2.0)]
    )
    add_page_break(doc)

    # ── SECTION 8: AGENT DESIGN PATTERN ───────────────────────────────────
    add_heading1(doc, '8.  Claude API Agent Pattern')

    add_heading2(doc, 'What Uses Claude (LLM)')
    add_bullet(doc, 'Generating vendor brief narrative from scorecard data')
    add_bullet(doc, 'Drafting 2–3 pushback response options (factual / neutral / escalation)')
    add_bullet(doc, 'Extracting structured action items from pasted meeting notes')
    add_bullet(doc, 'Generating meeting minutes from captured note items')
    add_bullet(doc, 'Generating leadership briefing card insights')
    add_bullet(doc, 'Change highlight summary text (Module C)')

    add_heading2(doc, 'What is Deterministic (No LLM)')
    add_bullet(doc, 'Slot ranking algorithm — hard/soft constraint scoring')
    add_bullet(doc, 'Score validation — range check, comment requirement rule')
    add_bullet(doc, 'Outlier detection — standard deviation calculation')
    add_bullet(doc, 'Score averaging and scorecard compilation')
    add_bullet(doc, 'Cycle-to-cycle score diff (simple delta)')
    add_bullet(doc, 'Alignment flag detection — spread ≥ 1.5 threshold')
    add_bullet(doc, 'Recurring issue detection — count query on issues table')
    add_bullet(doc, 'Workflow state transitions')

    add_info_box(doc, 'Design Principle:',
        'Deterministic logic first, AI second. This keeps the system explainable to a '
        'Shell executive audience and avoids hallucination in critical paths.')

    add_heading2(doc, 'Slot Ranking Algorithm (Module A)')
    add_standard_table(doc,
        ['Factor', 'Type', 'Rule'],
        [
            ['Organiser available', 'Hard constraint', 'Slot invalid if organiser blocked'],
            ['Exec sponsor available', 'Hard constraint', 'Slot invalid if exec sponsor blocked'],
            ['Max group attendance', 'Soft score', '(confirmed / total) × 100'],
            ['Conflict count', 'Penalty', '−10 per non-key attendee conflict'],
            ['Timezone suitability', 'Bonus', '+5 if within 09:00–17:00 local time for all key stakeholders'],
        ],
        [Inches(2.0), Inches(1.5), Inches(2.5)]
    )

    add_heading2(doc, 'Scorecard Validation Rules (Module B)')
    add_standard_table(doc,
        ['Rule', 'Type', 'Action'],
        [
            ['score < 1 or score > 5', 'ERROR', 'Reject submission — out of range'],
            ['score = 1 or 5, no comment', 'ERROR', 'Reject — comment required for extreme scores'],
            ['|score − group avg| > 1.5σ', 'WARNING', 'Flag as outlier — shown in compiled view'],
            ['Required category missing', 'ERROR', 'Reject — required field empty'],
        ],
        [Inches(2.2), Inches(1.0), Inches(2.8)]
    )
    add_page_break(doc)

    # ── SECTION 9: MODULE PLANS ────────────────────────────────────────────
    add_heading1(doc, '9.  Module-by-Module Implementation Plan')

    # Module A
    add_heading2(doc, 'Module A — Meeting Scheduling & Coordination')
    add_body(doc, 'Goal: Go from a blank cycle to a confirmed meeting invite with tracked RSVPs.')
    add_heading3(doc, 'Steps')
    steps_a = [
        '"Start New Cycle" → creates cycle record → CYCLE_CREATED state',
        'System loads attendees from previous cycle record',
        'Scheduling agent generates attendee refresh form (rendered as in-app modal)',
        'User reviews and dispatches → state moves to ATTENDEE_REFRESH_SENT',
        '"Simulate Responses" button populates mock responses from seed data',
        'Attendee list updated — new names added, old ones replaced',
        'MockCalendarService returns fixture availability for all confirmed attendees',
        'Deterministic slot ranker runs → top 3 ranked slots shown with attendance breakdown',
        'Organiser clicks "Approve This Slot" → invite draft generated',
        'Approval panel shows email preview → "Send Invite" → MockEmailService stores to outbox',
        'Confirmation tracker panel shows ACCEPTED / DECLINED / PENDING per attendee',
        'Auto-nudge message generated for non-responders (shown in Notifications)',
    ]
    for i, s in enumerate(steps_a, 1):
        add_bullet(doc, f'{i}. {s}')

    add_heading3(doc, 'UI Components')
    add_standard_table(doc,
        ['Component', 'Purpose'],
        [
            ['AttendeeRefreshPanel', 'Shows current list, refresh form, response status'],
            ['SlotRankingPanel', 'Three slot cards with attendance breakdown table'],
            ['InviteApprovalPanel', 'Email preview + approve / reject buttons'],
            ['ConfirmationTracker', 'Live RSVP table per attendee'],
        ],
        [Inches(2.5), Inches(3.5)]
    )

    # Module B
    add_heading2(doc, 'Module B — Scorecard Collection & Validation')
    add_body(doc, 'Goal: Collect, validate, flag outliers, and compile a final scorecard.')
    add_heading3(doc, 'Steps')
    steps_b = [
        'Triggered after meeting scheduled (can run in parallel)',
        'Scorecard request form generated — 5 categories, 1–5 scale, comment field per stakeholder',
        'Dispatch approval panel → MockEmailService sends to all stakeholders',
        'Reminder schedule runs automatically (visible in Notifications panel):',
        '  T−5 days → informational tone  |  T−2 days → deadline notice  |  T−day → escalation flag',
        'Each submission goes through deterministic validation service',
        'Invalid submissions trigger inline correction request back to submitter',
        'Scorecard status panel shows per-stakeholder submission progress',
        '"Compile Scorecard" → final table with averages + outlier flags → SCORECARD_COMPILED',
    ]
    for s in steps_b:
        add_bullet(doc, s)

    add_heading3(doc, 'UI Components')
    add_standard_table(doc,
        ['Component', 'Purpose'],
        [
            ['ScorecardDispatchPanel', 'Approval gate for sending scorecard requests'],
            ['SubmissionTracker', 'Per-stakeholder status table with reminder history'],
            ['CompiledScorecardTable', 'Final table with outlier badges and averages row'],
        ],
        [Inches(2.5), Inches(3.5)]
    )

    # Module C
    add_heading2(doc, 'Module C — Internal Alignment Call Support')
    add_body(doc, 'Goal: Prepare the internal team by surfacing changes and capturing actions.')
    add_heading3(doc, 'Steps')
    steps_c = [
        'Triggered after scorecard compiled',
        'Deterministic score diff engine compares current vs previous cycle — delta ≥ 1 point flagged',
        'Alignment flag engine: spread ≥ 1.5 points between stakeholders → prompt question generated',
        'Claude generates 3–5 bullet "What Changed" summary from diff data',
        'Face-off model panel shows current Shell/Vendor role-name mappings — inline editable',
        'Coordinator pastes meeting notes into text area',
        'Claude extracts structured action items: description, owner, due date',
        'Action items added to action log with status = OPEN',
    ]
    for s in steps_c:
        add_bullet(doc, s)

    add_heading3(doc, 'UI Components')
    add_standard_table(doc,
        ['Component', 'Purpose'],
        [
            ['ChangeHighlightsPanel', 'Bullet list of score deltas and new issues vs prior cycle'],
            ['AlignmentFlagsPanel', 'List of divergence flags with prompt questions for team to resolve'],
            ['FaceOffModelEditor', 'Numbered grid, inline editable Shell/Vendor name-role pairs'],
            ['NotesInputPanel', 'Paste area + "Extract Actions" button'],
            ['ActionLog (shared)', 'Table with filter by source module, status badges, owner, due date'],
        ],
        [Inches(2.5), Inches(3.5)]
    )

    # Module D
    add_heading2(doc, 'Module D — Vendor Prep Call Support')
    add_body(doc, 'Goal: Equip the Shell team with a vendor brief and structured pushback handling.')
    add_heading3(doc, 'Steps')
    steps_d = [
        'Triggered after internal alignment complete',
        'Claude generates vendor brief using tool calls: scorecard averages, comments, prior cycle, open actions',
        'Brief structured as: Overall Score, Category Ratings + Rationale, Key Concerns, Positive Areas',
        'Brief shown in approval panel — human reviews before it reaches vendor',
        'Vendor objection form: free text + category (DATA_DISPUTE / PROCESS_CONCERN / RESOURCE_CONSTRAINT / SCOPE_DISAGREEMENT / OTHER)',
        'Claude drafts 3 response options per pushback: Factual · Neutral Collaborative · Firm Escalation',
        'Items requiring legal/commercial review flagged — excluded from AI drafts',
        'Unresolved items stored in issues table with status = OPEN',
        'Face-off model update and notes/action capture (same as Module C)',
    ]
    for s in steps_d:
        add_bullet(doc, s)

    add_heading3(doc, 'UI Components')
    add_standard_table(doc,
        ['Component', 'Purpose'],
        [
            ['VendorBriefPanel', 'Structured brief card with approve / share button'],
            ['PushbackInput', 'Form to add objections with category selector'],
            ['PushbackResponseCards', 'Three option cards — factual / neutral / escalation — with select/edit'],
            ['UnresolvedItemTracker', 'Table with status badges, raised-by, date'],
        ],
        [Inches(2.5), Inches(3.5)]
    )

    # Module E
    add_heading2(doc, 'Module E — EGB/QBR Live Meeting Support')
    add_body(doc, 'Goal: Real-time capture during the meeting + automatic post-meeting artefact generation.')
    add_heading3(doc, 'Steps')
    steps_e = [
        'At meeting start, trend briefing card shown (from Module F analytics engine)',
        'Live capture panel: facilitator logs QUESTION / OBJECTION / DECISION / APPRECIATION / ACTION items',
        'Each item timestamped and tagged by type',
        'Alternate mode: paste full transcript → Claude parses into structured items',
        '"Generate Minutes" → Claude produces: metadata, exec summary, agenda summaries, key decisions, Q&A log, actions',
        'Minutes shown in approval panel',
        'On approval, action items merged with open items from Modules C and D',
        'Minutes available for copy-to-clipboard export',
    ]
    for s in steps_e:
        add_bullet(doc, s)

    add_heading3(doc, 'UI Components')
    add_standard_table(doc,
        ['Component', 'Purpose'],
        [
            ['MeetingBriefingCard', 'Pre-meeting trend summary — most improved, most concerning areas'],
            ['LiveCapturePanel', 'Type-selector + text input + timestamped running feed'],
            ['TranscriptInput', 'Paste-and-parse alternate mode for pre-written transcripts'],
            ['MeetingMinutesViewer', 'Structured minutes display + copy-to-clipboard button'],
            ['ActionLog (shared)', 'Final merged action log with all source modules'],
        ],
        [Inches(2.5), Inches(3.5)]
    )

    # Module F
    add_heading2(doc, 'Module F — Cross-Cycle Memory & Analytics Dashboard')
    add_body(doc, 'Goal: Persistent institutional memory, trend charts, recurring issue detection, leadership briefs.')
    add_heading3(doc, 'Steps')
    steps_f = [
        'Historical cycles auto-stored — 4 pre-seeded cycles + any new cycles',
        'Trend chart: per-vendor per-category line chart over 4 cycles (Recharts LineChart)',
        'Radar chart: current vs previous cycle overall vendor health (Recharts RadarChart)',
        'Cross-vendor bar chart: current cycle scores side-by-side (Recharts BarChart)',
        'Recurring issue detection (deterministic): issues table query — occurrences ≥ 2 AND status = OPEN',
        'Alert cards: "Delivery Quality flagged for 3 consecutive cycles — CoreSystems Ltd"',
        'Claude generates leadership briefing card: trajectory + recurring issues + prior commitments + focus areas',
        'Charts and brief accessible from main Analytics page and within each cycle workspace',
    ]
    for s in steps_f:
        add_bullet(doc, s)

    add_heading3(doc, 'UI Components')
    add_standard_table(doc,
        ['Component', 'Purpose'],
        [
            ['TrendLineChart', 'Per-vendor per-category scores over cycles'],
            ['RadarChart', 'Current vs prior cycle overall vendor health'],
            ['CrossVendorComparison', 'Bar chart — current cycle all vendors side-by-side'],
            ['RecurringIssueAlerts', 'Alert card list — flagged issues with occurrence count'],
            ['LeadershipBriefCard', '4-section card with "Generate" button — Claude powered'],
        ],
        [Inches(2.5), Inches(3.5)]
    )
    add_page_break(doc)

    # ── SECTION 10: FRONTEND ARCHITECTURE ─────────────────────────────────
    add_heading1(doc, '10.  Frontend Architecture')

    add_heading2(doc, 'Page Structure')
    add_standard_table(doc,
        ['Route', 'Page', 'Content'],
        [
            ['/', 'Dashboard', 'Active cycles overview · Start New Cycle CTA · Recent agent runs · Vendor quick-access'],
            ['/cycles/:cycleId', 'Cycle Workspace', 'Tabbed layout: Overview · Scheduling · Scorecard · Alignment · Vendor Prep · Meeting · Actions'],
            ['/analytics', 'Analytics Dashboard', 'Module F charts · Vendor selector · Leadership brief generator'],
        ],
        [Inches(1.5), Inches(1.5), Inches(3.2)]
    )

    add_heading2(doc, 'Cycle Workspace Tabs')
    add_standard_table(doc,
        ['Tab', 'Shows'],
        [
            ['Overview', 'Workflow progress bar · Summary status cards · Last agent run'],
            ['Scheduling', 'Module A — full scheduling workflow'],
            ['Scorecard', 'Module B — collection, validation, compilation'],
            ['Alignment', 'Module C — change highlights, flags, face-off, actions'],
            ['Vendor Prep', 'Module D — brief, pushback, responses, unresolved'],
            ['Meeting', 'Module E — live capture, minutes, merged actions'],
            ['Actions', 'Unified action log filtered across all modules for this cycle'],
        ],
        [Inches(1.5), Inches(4.5)]
    )

    add_heading2(doc, 'Shared UI Patterns')
    add_heading3(doc, 'Approval Panel — used before every "send" action')
    add_code_block(doc, """┌─────────────────────────────────────────────────────────┐
│  Agent summary:  Attendee refresh form ready            │
│  Preview:        [email/form content rendered here]     │
│  Recipients:     Alex (alex@shell.com), Priya, Marcus   │
│  [Approve & Send]          [Edit]          [Cancel]     │
└─────────────────────────────────────────────────────────┘""")

    add_heading3(doc, 'Workflow Progress Bar — always visible in cycle workspace')
    add_code_block(doc, """[Scheduling ✓] → [Scorecard ✓] → [Alignment ●] → [Vendor Prep] → [Meeting] → [Complete]""")

    add_body(doc, 'Agent Status Badge — shown on every module tab:  IDLE  ·  RUNNING  ·  AWAITING_APPROVAL  ·  COMPLETE  ·  FAILED')
    add_page_break(doc)

    # ── SECTION 11: KEY API ENDPOINTS ─────────────────────────────────────
    add_heading1(doc, '11.  Key API Endpoints')

    add_standard_table(doc,
        ['Method', 'Path', 'Purpose'],
        [
            ['POST', '/api/cycles', 'Create new governance cycle'],
            ['GET', '/api/cycles', 'List all cycles with workflow state'],
            ['GET', '/api/cycles/{id}', 'Cycle detail + current workflow state'],
            ['POST', '/api/cycles/{id}/scheduling/start', 'Trigger attendee refresh (Module A)'],
            ['POST', '/api/cycles/{id}/scheduling/simulate-responses', 'Demo: simulate attendee responses'],
            ['GET', '/api/cycles/{id}/scheduling/slots', 'Get ranked slot proposals'],
            ['POST', '/api/cycles/{id}/scheduling/approve-slot', 'Approve a slot — triggers invite draft'],
            ['POST', '/api/cycles/{id}/scorecard/send-request', 'Dispatch scorecard forms to stakeholders'],
            ['POST', '/api/cycles/{id}/scorecard/submit', 'Submit a score (mock form submission)'],
            ['POST', '/api/cycles/{id}/scorecard/compile', 'Compile final scorecard from valid inputs'],
            ['GET', '/api/cycles/{id}/scorecard/compiled', 'Get compiled scorecard with averages + flags'],
            ['GET', '/api/cycles/{id}/alignment/changes', 'Get score diffs vs prior cycle'],
            ['GET', '/api/cycles/{id}/alignment/flags', 'Get alignment divergence flags'],
            ['POST', '/api/cycles/{id}/alignment/extract-actions', 'Claude extracts actions from pasted notes'],
            ['POST', '/api/cycles/{id}/vendor-prep/generate-brief', 'Claude generates vendor brief'],
            ['POST', '/api/cycles/{id}/vendor-prep/pushback', 'Add pushback / objection item'],
            ['GET', '/api/cycles/{id}/vendor-prep/pushback/{pid}/responses', 'Get Claude-drafted response options'],
            ['POST', '/api/cycles/{id}/meeting/capture', 'Add live meeting note (typed or parsed)'],
            ['POST', '/api/cycles/{id}/meeting/generate-minutes', 'Claude generates meeting minutes'],
            ['GET', '/api/cycles/{id}/meeting/minutes', 'Get generated minutes for approval'],
            ['GET', '/api/analytics/vendors/{vid}/trends', 'Score trends over 4 cycles per category'],
            ['GET', '/api/analytics/recurring-issues', 'All recurring issue alerts across vendors'],
            ['POST', '/api/analytics/cycles/{id}/leadership-brief', 'Claude generates leadership briefing card'],
            ['GET', '/api/agent-runs', 'Full agent execution log (traceability)'],
        ],
        [Inches(0.6), Inches(2.8), Inches(2.6)]
    )
    add_page_break(doc)

    # ── SECTION 12: SEED DATA ──────────────────────────────────────────────
    add_heading1(doc, '12.  Seed Data Strategy')
    add_body(doc,
        'The seed data tells a deliberate story. Without narrative data, the analytics module '
        'is meaningless. Each vendor has a distinct trajectory.')

    add_heading2(doc, 'Vendor Score Trajectories (Q1→Q2→Q3→Q4)')
    add_standard_table(doc,
        ['Category', 'NovaTech (Improving)', 'CoreSystems (Declining)', 'Meridian (Stable)'],
        [
            ['Delivery Quality',   '3 → 3 → 4 → 4', '4 → 3 → 3 → 2', '3 → 3 → 3 → 3'],
            ['SLA Compliance',     '2 → 3 → 3 → 4', '3 → 3 → 2 → 2', '4 → 4 → 3 → 4'],
            ['Innovation',         '3 → 3 → 4 → 5', '3 → 2 → 2 → 2', '3 → 3 → 3 → 3'],
            ['Communication',      '3 → 4 → 4 → 4', '4 → 3 → 3 → 2', '3 → 3 → 4 → 3'],
            ['Value for Money',    '3 → 3 → 3 → 4', '4 → 4 → 3 → 3', '3 → 3 → 3 → 3'],
        ],
        [Inches(1.8), Inches(1.8), Inches(1.8), Inches(1.8)]
    )

    add_heading2(doc, 'Pre-Seeded Recurring Issues')
    add_standard_table(doc,
        ['Vendor', 'Issue', 'Flagged In', 'Status'],
        [
            ['CoreSystems Ltd', 'Delivery Quality consistently below SLA threshold', 'Q2, Q3, Q4', 'OPEN'],
            ['CoreSystems Ltd', 'Delayed invoice submissions', 'Q3, Q4', 'OPEN'],
            ['NovaTech Services', 'Innovation KPIs not aligned to contract commitments', 'Q2, Q3', 'RESOLVED (Q4)'],
        ],
        [Inches(1.5), Inches(2.8), Inches(1.2), Inches(0.8)]
    )

    add_heading2(doc, 'Other Seed Requirements')
    add_bullet(doc, '8–10 stakeholders across Shell VMO, IDT Operations, and vendor teams')
    add_bullet(doc, '4 historical EGB cycles (Q1–Q4 previous year) with realistic score variations')
    add_bullet(doc, 'Pre-seeded action items — mix of OPEN and CLOSED — across historical cycles')
    add_bullet(doc, 'At least 2 pre-seeded objection items per vendor for pushback demo')
    add_bullet(doc, 'Face-off model seeded with 6 Shell roles and 4 vendor roles per cycle')
    add_page_break(doc)

    # ── SECTION 13: BUILD SEQUENCE ─────────────────────────────────────────
    add_heading1(doc, '13.  Build Sequence — 5-Day Sprint')

    add_phase_box(doc, 'Day 1', 'Foundation', [
        'Project scaffolding — Vite + FastAPI + SQLite',
        'Database setup — all 13 tables via Alembic migrations',
        'Seed data script — 4 cycles, 3 vendors, realistic scores',
        'Workflow engine — state machine class with enforcement rules',
        'Base agent pattern + agent_runs logging',
        'Mock services — calendar, email, forms, notifications',
        'Standard AgentResponse Pydantic schema (contract)',
        'Frontend: routing, layout shell, Tailwind config, shadcn/ui setup',
    ])

    add_phase_box(doc, 'Day 2', 'Module A + B', [
        'Scheduling agent — attendee refresh, slot ranking (deterministic), invite draft',
        'Scheduling UI — AttendeeRefreshPanel, SlotRankingPanel, InviteApprovalPanel, ConfirmationTracker',
        '"Simulate Responses" demo button for Module A',
        'Validation service — deterministic rules only (no LLM)',
        'Scorecard agent — compilation, averaging',
        'Scorecard UI — DispatchPanel, SubmissionTracker, CompiledScorecardTable',
        '"Simulate Submissions" demo button for Module B',
    ])

    add_phase_box(doc, 'Day 3', 'Module C + D', [
        'Score diff engine — deterministic comparison (score_diff.py)',
        'Alignment flag engine — deterministic spread check',
        'Claude integration — action item extraction from notes (Module C)',
        'Claude integration — vendor brief generation (Module D)',
        'Claude integration — pushback response drafting 3 options (Module D)',
        'Alignment UI — ChangeHighlights, AlignmentFlags, FaceOffModelEditor, NotesInput',
        'Vendor Prep UI — VendorBriefPanel, PushbackInput, ResponseCards, UnresolvedTracker',
        'Shared ActionLog component',
    ])

    add_phase_box(doc, 'Day 4', 'Module E + F', [
        'Live capture panel with all 5 note types',
        'Transcript paste + Claude parsing mode',
        'Claude meeting minutes generation',
        'Trend charts — Recharts LineChart, RadarChart, BarChart',
        'Recurring issue detection — deterministic DB query',
        'Claude leadership brief generation',
        'Meeting UI — BriefingCard, LiveCapturePanel, TranscriptInput, MinutesViewer',
        'Analytics page — 3 chart types, RecurringIssueAlerts, LeadershipBriefCard',
    ])

    add_phase_box(doc, 'Day 5', 'Integration + Demo', [
        'End-to-end workflow state machine wiring — all 6 modules connected',
        '8-step demo narrative walkthrough — test and fix',
        'Approval panels verified for all send actions',
        'Agent execution trace log page',
        'Error handling — LLM retry, graceful fallback text, empty states',
        'Demo data review — story reads correctly across all vendors',
        'Performance check — all responses < 5 seconds',
        'Deploy to public URL (FastAPI → Render/Fly.io · React → Vercel)',
    ])

    add_page_break(doc)

    # ── SECTION 14: DEMO FLOW ──────────────────────────────────────────────
    add_heading1(doc, '14.  Demo Flow Alignment — 8-Step Narrative')
    add_body(doc,
        'Every step is triggerable by a single button or form action. '
        'No manual data entry is required during the 30-minute demo.')
    add_standard_table(doc,
        ['Step', 'Module', 'Trigger', 'What Is Shown'],
        [
            ['1', 'A', '"Start New Cycle" button', 'Attendee refresh form dispatched — 9 stakeholders notified'],
            ['2', 'A', '"Simulate Responses" button', 'Agent ranks top 3 slots with attendance breakdown — organiser approves'],
            ['3', 'B', '"Send Scorecard Request"', 'Dispatch approval panel + full 3-tier reminder schedule visible'],
            ['4', 'B', '"Simulate Submissions"', 'Outlier flagged in compiled scorecard — averages calculated'],
            ['5', 'C', 'Auto-triggered on compile', 'Change highlights vs prior cycle — one alignment flag raised'],
            ['6', 'D', '"Generate Vendor Brief"', 'Brief generated → pushback entered → 3 response drafts shown'],
            ['7', 'E', '"Start Meeting" button', 'Live notes captured → minutes auto-generated → merged action log'],
            ['8', 'F', 'Navigate to Analytics', 'Recurring issue alert fires → leadership brief card generated'],
        ],
        [Inches(0.4), Inches(0.6), Inches(1.8), Inches(3.2)]
    )
    add_page_break(doc)

    # ── SECTION 15: KEY DECISIONS ──────────────────────────────────────────
    add_heading1(doc, '15.  Key Technical & Design Decisions')
    add_standard_table(doc,
        ['Decision', 'Rationale'],
        [
            ['SQLite over Postgres', 'Zero-config, file-based, demo-reset friendly. Schema is Postgres-compatible for production migration.'],
            ['Tool-calling agents over prompt-only', 'Better control, lower hallucination, easier debugging, more production-like behaviour.'],
            ['Deterministic logic for validation and ranking', 'Explainable to a Shell executive audience. No AI black-box in critical governance paths.'],
            ['Mock services behind clean interfaces', 'Swap to Outlook/Teams/SharePoint API later without rewriting any agent code.'],
            ['Single JSON contract (AgentResponse)', 'Frontend never guesses shape. Stable even as AI outputs evolve. Enables reliable approval flows.'],
            ['Human approval before every "send" action', 'Builds client trust. Legally important for governance context. AI assists, humans decide.'],
            ['Workflow state machine', 'Prevents impossible combinations — cannot generate minutes without notes, cannot compile without submissions.'],
            ['Seed data tells a story', 'Analytics only work if data has deliberate trajectories and recurring issues. Filler data breaks the demo narrative.'],
            ['Zustand + TanStack Query', 'Zustand for UI state (active cycle, approval modals). TanStack Query for server state with automatic background refresh.'],
            ['Recharts for visualisation', 'Tailwind-compatible, lightweight, supports all 3 required chart types (Line, Radar, Bar).'],
        ],
        [Inches(2.2), Inches(4.0)]
    )

    # ── NON-FUNCTIONAL REQUIREMENTS ───────────────────────────────────────
    add_heading2(doc, 'Non-Functional Requirements')
    add_standard_table(doc,
        ['Category', 'Requirement'],
        [
            ['Performance', 'All agent responses render within 5 seconds in demo environment'],
            ['Reliability', 'Demo flows must be reproducible with no manual resets required between runs'],
            ['Mock data fidelity', 'All demo data uses realistic energy-sector vendor names, categories, and scores'],
            ['LLM output quality', 'All LLM-generated outputs reviewed and locked before demo day'],
            ['Accessibility', 'UI must be usable on a standard laptop browser at 1280×800 resolution'],
            ['Portability', 'Deployed on a single public URL accessible without VPN'],
        ],
        [Inches(1.8), Inches(4.4)]
    )

    # ── RISKS ─────────────────────────────────────────────────────────────
    add_heading2(doc, 'Risks & Mitigations')
    add_standard_table(doc,
        ['Risk', 'Mitigation'],
        [
            ['LLM hallucination in generated minutes or briefs', 'All LLM outputs pre-reviewed and locked to a fixed demo transcript; live generation used for non-critical fields only'],
            ['Demo data feels unrealistic', 'Seed data reviewed by architect and domain-aware team member before Day 4'],
            ['Integration complexity between modules', 'All modules share a single JSON contract; orchestration layer enforces schema'],
            ['Scope creep from team', 'Architect holds scope freeze after Day 1; any new feature goes into backlog, not the sprint'],
        ],
        [Inches(2.5), Inches(3.7)]
    )

    # ── GLOSSARY ──────────────────────────────────────────────────────────
    add_page_break(doc)
    add_heading1(doc, 'Appendix — Glossary')
    add_standard_table(doc,
        ['Term', 'Definition'],
        [
            ['EGB', 'Executive Governance Board — Shell\'s senior vendor governance forum'],
            ['QBR', 'Quarterly Business Review — operational vendor performance review cycle'],
            ['VMO', 'Vendor Management Office — Shell team responsible for vendor governance'],
            ['IDT', 'Information & Digital Technology — Shell\'s IT function'],
            ['Face-off model', 'Structured table mapping Shell and vendor roles/names to governance responsibilities'],
            ['Cross-cycle memory', 'Persistent storage of historical scorecard data, actions, and decisions across governance cycles'],
            ['Agentic AI', 'An AI system that takes multi-step actions using tools and memory to complete a goal, with human oversight checkpoints'],
            ['Workflow state', 'The current stage of a governance cycle in the state machine — drives which modules are active'],
            ['Tool-calling', 'Claude API pattern where the LLM decides which Python functions to call rather than generating raw text answers'],
        ],
        [Inches(1.8), Inches(4.4)]
    )

    # ── SAVE ──────────────────────────────────────────────────────────────
    output_path = r'c:\Users\AK115384\Desktop\VenderPulse\dev artifacts\VendorPulse_MVP_Dev_Plan_v2.docx'
    doc.save(output_path)
    print(f'Document saved: {output_path}')
    return output_path


if __name__ == '__main__':
    build_document()
